#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os, sys, re, glob
import pandas as pd
import time
import argparse
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from xml.dom import minidom

import smtplib
from email.mime.text import MIMEText

path2script = os.path.realpath(os.path.dirname(__file__))

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

import numpy as np

import  utils.parseRunParameters  as parseRunParameters
from utils.wreportOneSample import multiRunCommentsAndNotes1Sample, rnaRunCommentsAndNotes1Sample
from utils.utils import generate_DME_link, replace_string_hyperlink 

def replace_string(doc, before, after):
    for p in doc.paragraphs:
        if before in p.text:
            inline = p.runs

            for i in range(len(inline)):
                if before in inline[i].text:
                    text = inline[i].text.replace(before, str(after))
                    inline[i].text = text
                    #return text

    return 1

def bulletEntry(rc, header, description):
    paragraph = rc.insert_paragraph_before(style='List Bullet')
    runner = paragraph.add_run(header)
    runner.bold = True
    runner = paragraph.add_run(f" - {description}")
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.left_indent = Inches(0.25)

def getCellRangerVersion(values, de, cellranger, pipeline):
    pattern = re.compile(r'\d+\.\d+\.\d+')
    if len(pattern.findall(cellranger)) == 1:
        values['version'] = pattern.findall(cellranger)[0]
        if pipeline.lower() == 'multiome':
            values['version'] = "ARC " + values['version']
        elif pipeline.lower() == 'atac':
            values['version'] = "ATAC " + str(list(de[u'Pipeline version'])[0].split('-')[-1])
        elif pipeline.lower() == 'atac_1.2.0':
            values['version'] = "ATAC " + str(list(de[u'cellranger-atac_version'])[0])
    else:
        values['version'] = "UNKNOWN"
    return(values)

def getCellRangerCommand(pipeline):
    if pipeline.lower() == "multiome":
        return("cellranger-arc count --id=sample_ID --reference=reference_dir --libraries=libraries.csv")
    elif pipeline.lower() == "rna":
        return("cellranger count --id=sample_ID --sample=sample_ID --fastqs=fastq_path --transcriptome=reference_dir")
    elif pipeline.lower() == "fb":
        return("cellranger count --id=sample_ID --transcriptome=reference_dir --libraries=libraries.csv --feature-ref=features.csv")
    elif pipeline.lower() == "atac" or pipeline.lower() == "atac_1.2.0":
        return("cellranger count --id=sample_ID --fastqs=fastq_path --reference=reference_dir")
    elif pipeline.lower() == "vdj":
        return("cellranger vdj --id=sample_ID --reference=reference_dir --fastqs=fastq_path")
    elif pipeline.lower() == "multi":
        return("cellranger multi --id=sample_ID --csv=sample_ID.csv")
    elif pipeline.lower() == "pipseq":
        return("pipseeker full --fastq fastq_path --output-path output_folder --star-index-path path2star_index")
    else:
        return("UNKNOWN")

def getReference(species, pipeline):
    if pipeline.lower() == "multiome" or pipeline.lower() == "atac":
        if "HG38" in species.upper():
            return("hg38 - 2020-A-2.0.0")
        elif "MM10" in species.upper():
            return("mm10 - 2020-A-2.0.0")
    elif any([pipeline.lower() == i for i in ["rna", "fb"]]):
        if "HG38" in species.upper():
            return("hg38 - 2020-A")
        elif "MM10" in species.upper():
            return("mm10 - 2020-A")
    elif pipeline.lower() == "atac_1.2.0":
        if "HG38" in species.upper():
            return("hg38 - 1.2.0")
        elif "MM10" in species.upper():
            return("mm10 - 1.2.0")
    elif pipeline.lower() == "vdj":
        if "HG38" in species.upper():
            return("GRCh38 - 5.0.0")
        elif "MM10" in species.upper():
            return("GRCm38 - 5.0.0")
    return("UNKNOWN")

def getLibraryProtocol(pipeline):
    if pipeline.lower() == "multiome":
        return("Single Cell Multiome Gene Expression and ATAC")
    elif pipeline.lower() == "rna":
        return("Single Cell Gene Expression")
    elif pipeline.lower() == 'fb':
        return("Single Cell Gene Expression with Feature Barcode for Surface protein")
    elif pipeline.lower() == 'vdj':
        return("Single Cell VDJ")
    elif pipeline.lower() == "atac" or pipeline.lower() == "atac_1.2.0":
        return("Single Cell ATAC")
    else:
        return("UNKNOWN")

def parseXml(xmlRunParametersPath, xmlRunInfoPath, instrument):
    RTAVersion = "Unknown"
    flowcell = "Unknown"
    workFlowTypeValue = "Unknown"
    flowcellModeValue = "Unknown"
    chemistryValue = "Unknown"
    chemistryVersionValue = "Unknown"
    applicationValue = "Unknown"
    xmldoc = minidom.parse(xmlRunParametersPath)
    tagRTA = ""
    if instrument == "NovaSeq 6000" or instrument == "NextSeq 2000" or instrument == "iSeq":
        tagRTA = "RtaVersion"
    elif instrument == "HiSeq" or  instrument == "NextSeq 550" or instrument == "MiSeq":
        tagRTA = "RTAVersion"
    elif instrument == "NovaSeq Xplus":
        RTAVersion = "RTA4"
    if tagRTA:
        cyc = xmldoc.getElementsByTagName(tagRTA)
    else:
        cyc = []
    if len(cyc) > 0:
        RTAVersion = cyc[0].childNodes[0].nodeValue
    workFlowType = xmldoc.getElementsByTagName('WorkflowType')
    if len(workFlowType) > 0:
        workFlowTypeValue = workFlowType[0].childNodes[0].nodeValue
    flowcellMode = xmldoc.getElementsByTagName('FlowCellMode')
    if len(flowcellMode) > 0:
        flowcellModeValue = flowcellMode[0].childNodes[0].nodeValue
    chemistry = xmldoc.getElementsByTagName('Chemistry')
    if len(chemistry) > 0:
        chemistryValue = chemistry[0].childNodes[0].nodeValue
    chemistryVersion = xmldoc.getElementsByTagName('ChemistryVersion')
    if len(chemistryVersion) > 0:
        chemistryVersionValue = chemistryVersion[0].childNodes[0].nodeValue
    application = xmldoc.getElementsByTagName('Application')
    if len(application) > 0:
        applicaitonValue = application[0].childNodes[0].nodeValue
    consumableInfo = xmldoc.getElementsByTagName('ConsumableInfo')
    if len(consumableInfo) > 0:
        for consumableInfoIndex in range(len(consumableInfo)):
            if consumableInfo[consumableInfoIndex].getElementsByTagName('Type')[0].childNodes[0].nodeValue == 'FlowCell':
                if len(consumableInfo[consumableInfoIndex].getElementsByTagName('Name')) > 0:
                    flowcellModeValue = consumableInfo[consumableInfoIndex].getElementsByTagName('Name')[0].childNodes[0].nodeValue
                else:
                    flowcellModeValue = consumableInfo[consumableInfoIndex].getElementsByTagName('Mode')[0].childNodes[0].nodeValue
            else:
                continue
    xmldoc = minidom.parse(xmlRunInfoPath)
    runInfoFlowcell = xmldoc.getElementsByTagName('Flowcell')
    if len(runInfoFlowcell) > 0:
        flowcell = runInfoFlowcell[0].childNodes[0].nodeValue
    return (RTAVersion, flowcell, workFlowTypeValue, flowcellModeValue, chemistryValue, chemistryVersionValue)

def getRunPath(sRunName):
    instrument = ''
    runPath = ''
    if sRunName.group(2)[0] == "A":
        instrument = "NovaSeq 6000"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_NovaSeq", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_NovaSeq")
    elif sRunName.group(2)[0] == "L":
        instrument = "NovaSeq Xplus"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_Xplus", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_Xplus")
    elif sRunName.group(2)[0] == "J" or sRunName.group(2)[0] == "D":
        instrument = "HiSeq"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_HiSeq", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_HiSeq")
    elif sRunName.group(2)[0] == "N":
        instrument = "NextSeq 550"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_NextSeq", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_NextSeq")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_NextSeq", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_NextSeq")
    elif sRunName.group(2)[0] == "V":
        instrument = "NextSeq 2000"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_NextSeq", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_NextSeq")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_NextSeq", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_NextSeq")
    elif sRunName.group(2)[0] == "M":
        instrument = "MiSeq"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_MiSeq", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_MiSeq")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_MiSeq", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_MiSeq")
    elif sRunName.group(2)[0] == "F":
        instrument = "iSeq"
        if os.path.isdir(os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_iSeq", sRunName.group(0))):
            runPath = os.path.join('/is2/projects/CCR-SF/scratch/illumina', "RawData_iSeq")
        elif os.path.isdir(os.path.join('/mnt/ccrsf-raw/illumina', "RawData_iSeq", sRunName.group(0))):
            runPath = os.path.join('/mnt/ccrsf-raw/illumina', "RawData_iSeq")
    return(instrument, runPath)

def getRunInfo(run_names):
    runInfo = dict()
    #for run in run_names:
    #    #for filename in glob.glob(f'/is2/projects/CCR-SF/scratch/illumina/RawData*/{run}/*unParameters.xml'):
    #    for filename in glob.glob(f'/mnt/ccrsf-raw/illumina/RawData*/{run}/*unParameters.xml'):
    #        runInfo[run] = parseRunParameters.parseXML(filename)
    oRunName = re.compile("(\d{6,})_([A-Z0-9]+)_(\d+)_([-A-Z0-9]+)")
    for runName in run_names:
        sRunName = oRunName.search(runName)
        if sRunName:
            (instrument, runPath) = getRunPath(sRunName)
            xmlRunParametersPath = glob.glob(os.path.join(runPath, runName, "[rR]unParameters.xml" ))[0]
            xmlRunInfoPath = glob.glob(os.path.join(runPath, runName, "[rR]unInfo.xml" ))[0]
            (RTAVersion, flowcell, workFlowType, flowcellMode, chemistry, chemistryVersion) = parseXml(xmlRunParametersPath, xmlRunInfoPath, instrument)
            runInfo[runName] = parseRunParameters.parseXML(xmlRunParametersPath)
            if runInfo[runName]['platform'] == 'No Info':
                runInfo[runName]['platform'] = instrument
            if runInfo[runName]['flowcell'] == 'No Info':
                runInfo[runName]['flowcell'] = flowcell
            if runInfo[runName]['workflow'] == 'No Info':
                runInfo[runName]['workflow'] = workFlowType
            if runInfo[runName]['chemistry'] == 'No Info':
                runInfo[runName]['chemistry'] = flowcellMode
            if runInfo[runName]['rta'] == 'No Info':
                runInfo[runName]['rta'] = RTAVersion
        else:
            sys.stdout.write(f'Can not identify the run name from {runName}\n')
    return(runInfo)

def finalTableSetup(p, headers):
    p.style = 'Table Grid'
    for i in range(len(headers)-len(p.columns)):
        p.add_column(Inches(2))

    p.rows[0].height = Inches(0.8)
    for i in range(len(headers)):
        p.cell(0,i).paragraphs[0].text = ""
        runner = p.cell(0,i).paragraphs[0].add_run(headers[i])
        runner.bold=True
        runner.italic=True
        p.cell(0,i).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    while len(p.rows) > 1:
        row = p.rows[1]
        row._element.getparent().remove(row._element)

    while len(p.columns) > len(headers):
        col = p.table.columns[-1]
        grid = p._tbl.find("w:tblGrid", p._tbl.nsmap)
        for cell in col.cells:
            cell._element.getparent().remove(cell._element)
        # Delete column reference.
        col_elem = grid[-1]
        grid.remove(col_elem)

def multiomeRunCommentsAndNotes(doc, values, de):
    barcodeq30 = str(round(float(min(list(de[u'ATAC Q30 bases in barcode']) + list(de[u'GEX Q30 bases in barcode'])))*100, 2))
    umiq30 = str(round(float(min(list(de[u'GEX Q30 bases in UMI'])))*100, 2))
    readq30 = str(round(float(min([min(de[i]) for i in de.columns if 'Q30 bases in read' in i]))*100, 2))
    cellmin = "{:,}".format(int(min(list(de[u'Estimated number of cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'Estimated number of cells']))))

    rnarangemin = str(int(min(list(de[u'GEX Sequenced read pairs']))/1000000))
    atacrangemin = str(int(min(list(de[u'ATAC Sequenced read pairs']))/1000000))
    atacreadmin = "{:,}".format(int(min(list(de[u'ATAC Mean raw read pairs per cell']))))
    atacreadmax ="{:,}".format(int(max(list(de[u'ATAC Mean raw read pairs per cell']))))
    rnareadmin = "{:,}".format(int(min(list(de[u'GEX Mean raw reads per cell']))))
    rnareadmax ="{:,}".format(int(max(list(de[u'GEX Mean raw reads per cell']))))
    medgenemin = "{:,}".format(int(min(list(de[u'GEX Median genes per cell']))))
    medgenemax ="{:,}".format(int(max(list(de[u'GEX Median genes per cell']))))
    medfragmin = "{:,}".format(int(min(list(de[u'ATAC Median high-quality fragments per cell']))))
    medfragmax ="{:,}".format(int(max(list(de[u'ATAC Median high-quality fragments per cell']))))
    minfragtss = str(round(float(min(list(de[u'ATAC Fraction of high-quality fragments overlapping TSS'])))*100, 2))
    minfragpeak = str(round(float(min(list(de[u'ATAC Fraction of high-quality fragments overlapping peaks'])))*100, 2))

    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("Multiome Gene Expression")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on ")
    runner = rc.add_run(f"a {values['platform'].replace('_', ' ')} run")
    runner.bold = True
    runner = rc.add_run(f" and {values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("Multiome ATAC")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on ")
    runner = rc.add_run(f"a {values['platform'].replace('_', ' ')} run.")
    runner.bold = True
    runner = rc.add_run(f" All GEX samples have sequencing yields of more than {rnarangemin} million read pairs per sample and all ATAC samples have sequencing yields of more than {atacrangemin} million read pairs per sample. ")
    runner = rc.add_run(f"Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {readq30}% of bases in the read have Q30 or above. More than {umiq30}% of bases in the UMI have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax}. The mean raw reads per cell for GEX ranges from {rnareadmin} to {rnareadmax}. The mean raw read pairs per cell for ATAC ranges from {atacreadmin} to {atacreadmax}. Cells with extremely low number of UMI counts were filtered out. Median genes found per cell ranges from {medgenemin} to {medgenemax}. The median fragments per cell ranges from {medfragmin} to {medfragmax}. The fraction of fragments (that passed all filters) overlapping TSS is above {minfragtss}% and overlapping peaks is above {minfragpeak}%.")

    notesParagraph = doc.paragraphs[31]
    bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
    bulletEntry(notesParagraph, "Median High-Quality Fragments per Cell", "The median number of high-quality fragments per cell barcode.")
    bulletEntry(notesParagraph, "Confidently Mapped Read Pairs", "Fraction of sequenced read pairs with mapping quality > 30.")
    bulletEntry(notesParagraph, "Median Genes per Cell", "The median number of genes detected per cell barcode.")
    bulletEntry(notesParagraph, "Total Genes Detected", "The number of genes with at least one UMI count in any cell barcode.")
    bulletEntry(notesParagraph, "Median UMI Counts per Cell", "The median number of UMI counts per cell barcode.")

def multiomeFinalTable(doc, de):
    headers = ["Sample", "Estimated Number of Cells", "Median High-Quality Fragments per Cell", "ATAC Confidently Mapped Read Pairs", "Median Genes per Cell", "Total Genes Detected", "Median UMI Counts per Cell"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        p.cell(i+1,0).text = str(de[u"Sample ID"][i])
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,1).text = "{:,}".format(int(de[u"Estimated number of cells"][i]))
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,2).text = "{:,}".format(int(de[u"ATAC Median high-quality fragments per cell"][i]))
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,3).text = str(round(float(de[u"ATAC Confidently mapped read pairs"][i])*100, 2)) + "%"
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,4).text = "{:,}".format(int(de[u"GEX Median genes per cell"][i]))
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,5).text = "{:,}".format(int(de[u"GEX Total genes detected"][i]))
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,6).text = "{:,}".format(int(de[u"GEX Median UMI counts per cell"][i]))
        p.cell(i+1,6).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,6).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,6).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

def rnaRunCommentsAndNotesWithAntibody(doc, values, de, library_type):
    """
    """
    rc = doc.paragraphs[17] 
    lib_type_str = f"{library_type}:"
    lib_type_header = "Antibody"
    if library_type == "Custom":
        lib_type_header = "Feature"
    if any(lib_type_str in i for i in de):
        rangemin = str(int(min(list(de[f'{lib_type_str} Number of Reads']))/1000000))
        readmin = "{:,}".format(int(min(list(de[f'{lib_type_str} Mean Reads per Cell']))))
        readmax ="{:,}".format(int(max(list(de[f'{lib_type_str} Mean Reads per Cell']))))
        # Fraction of reads that contain a recognized antibody barcode
        min_frac_reads = "{:.2%}".format(min(list(de[f'{lib_type_str} Fraction {lib_type_header} Reads'])))
        max_frac_reads ="{:.2%}".format(max(list(de[f'{lib_type_str} Fraction {lib_type_header} Reads'])))
        # Number of antibody reads usable divided by the number of cell-associated barcodes.
        min_frac_reads_usable = "{:.2%}".format(min(list(de[f'{lib_type_str} Fraction {lib_type_header} Reads Usable'])))
        max_frac_reads_usable ="{:.2%}".format(max(list(de[f'{lib_type_str} Fraction {lib_type_header} Reads Usable'])))
        # Fraction of reads that contain a recognized antibody barcode, a valid UMI, and a cell-associated barcode.
        minpct_read_in_cell = "{:.2%}".format(min(list(de[f'{lib_type_str} {lib_type_header} Reads in Cells'])))
        maxpct_read_in_cell ="{:.2%}".format(max(list(de[f'{lib_type_str} {lib_type_header} Reads in Cells'])))

        runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
        runner = rc.add_run(f"{library_type.lower()}")
        runner.bold = True
        runner = rc.add_run(" libraries were sequenced on a ")
        runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
        runner.bold = True
        runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
        runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run.")
        runner.bold = True
        runner = rc.add_run(f" Mean reads per cell ranges from {readmin} to {readmax}. The fraction of reads that contain a recognized barcode ranges from {min_frac_reads} to {max_frac_reads}. The fraction of reads that contain a recognized barcode, a valid UMI, and a cell-associated barcode ranges from {min_frac_reads_usable} to {max_frac_reads_usable}. Among the reads that contain a recognized barcode, a valid UMI, and a valid barcode, the fraction associated with cell-containing partitions ranges from {minpct_read_in_cell} to {maxpct_read_in_cell}.\n\n")

def rnaRunCommentsAndNotes(doc, values, de):
    barcodeq30 = str(float(min(list(de[u'Q30 Bases in Barcode'])))*100)
    umiq30 = str(float(min(list(de[u'Q30 Bases in UMI'])))*100)
    cellmin = "{:,}".format(int(min(list(de[u'Estimated Number of Cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'Estimated Number of Cells']))))

    if sum(['RNA Read' in i for i in de.columns ]) == 1:
        rnareadq30 = str(float(min(list(de[u'Q30 Bases in RNA Read'])))*100)
    else:
        rnareadq30 = str(round(float(min([min(de[i]) for i in de.columns if 'Q30 Bases in RNA Read' in i]))*100, 2))

    rangemin = str(int(min(list(de[u'Number of Reads']))/1000000))
    rangemax = str(int(max(list(de[u'Number of Reads']))/1000000))
    readmin = "{:,}".format(int(min(list(de[u'Mean Reads per Cell']))))
    readmax ="{:,}".format(int(max(list(de[u'Mean Reads per Cell']))))
    medgenemin = "{:,}".format(int(min(list(de[u'Median Genes per Cell']))))
    medgenemax ="{:,}".format(int(max(list(de[u'Median Genes per Cell']))))
    totgenemin = "{:,}".format(int(min(list(de[u'Total Genes Detected']))))
    totgenemax ="{:,}".format(int(max(list(de[u'Total Genes Detected']))))

    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("Gene Expression")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {rnareadq30}% of bases in the read have Q30 or above. More than {umiq30}% of bases in the UMI have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax} and mean reads per cell ranges from {readmin} to {readmax}. Cells with extremely low number of UMI counts were filtered out. Median genes found per cell ranges from {medgenemin} to {medgenemax} and the total number of genes detected ranges from {totgenemin} to {totgenemax}.\n\n")

    rnaRunCommentsAndNotesWithAntibody(doc, values, de, "Antibody")
    rnaRunCommentsAndNotesWithAntibody(doc, values, de, "Custom") 

    notesParagraph = doc.paragraphs[31]
    bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
    bulletEntry(notesParagraph, "Mean Reads per Cell", "The total number of sequenced reads divided by the estimated number of cells.")
    bulletEntry(notesParagraph, "Median Genes per Cell", "The median number of genes detected per cell barcode.")
    bulletEntry(notesParagraph, "Median UMI Counts per Cell", "The median number of UMI counts per cell barcode.")
    bulletEntry(notesParagraph, "Number of Reads", "Total number of sequenced reads.")
    bulletEntry(notesParagraph, "Reads Mapped Confidently to Transcriptome", "Fraction of reads that mapped to a unique gene in the transcriptome with a high mapping quality score as reported by the aligner.")
    bulletEntry(notesParagraph, "Valid Barcodes", "Fraction of reads with cell-­barcodes that match the whitelist.")
    bulletEntry(notesParagraph, "Sequencing Saturation", "The fraction of reads originating from an already-­observed UMI. This is a function of library complexity and sequencing depth. More specifically, this is the fraction of confidently mapped, valid cell-­barcode, valid UMI reads that had a non-­unique (cell-­barcode, UMI, gene).")

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def replace_workflow_image(doc):
    """
    Replace or delete workflow image (if no workflow image is provided)
    """
    # Iterate over the paragraphs in the document
    for para in doc.paragraphs:
        # Iterate over the runs in the paragraph
        for run in para.runs:
            # Check if the run contains an image
            if "pic:pic" in run._element.xml:
                print(run._element)
                # Remove the run
                run._element.getparent().remove(run._element)
                break
        else:
            continue
        if args.workflow is not None:
            # Add workflow image
            para.add_run().add_picture(args.workflow)
        break

def pipseqRunCommentsAndNotes(doc, values, de):
    if 'sensitivity' in de.iloc[0]['sensitivity_level']:
        de = de[de['sensitivity_level'] == 'sensitivity_3']
        values['count'] = int(int(values['count'])/5) 
    cellmin = "{:,}".format(int(min(list(de[u'num_cell_barcodes']))))
    cellmax = "{:,}".format(int(max(list(de[u'num_cell_barcodes']))))

    #if sum(['RNA Read' in i for i in de.columns ]) == 1:
    #    rnareadq30 = str(float(min(list(de[u'Q30 Bases in RNA Read'])))*100)
    #else:
    #    rnareadq30 = str(round(float(min([min(de[i]) for i in de.columns if 'Q30 Bases in RNA Read' in i]))*100, 2))

    rangemin = str(int(min(list(de[u'num_input_reads']))/1000000))
    rangemax = str(int(max(list(de[u'num_input_reads']))/1000000))
    readmin = "{:,}".format(int(min(list(de[u'mean_reads_per_cell']))))
    readmax ="{:,}".format(int(max(list(de[u'mean_reads_per_cell']))))
    medgenemin = "{:,}".format(int(min(list(de[u'median_genes_in_cells']))))
    medgenemax ="{:,}".format(int(max(list(de[u'median_genes_in_cells']))))
    totgenemin = "{:,}".format(int(min(list(de[u'num_genes_expressed_in_cells']))))
    totgenemax ="{:,}".format(int(max(list(de[u'num_genes_expressed_in_cells']))))

    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} PIPseq Single Cell ")
    runner = rc.add_run("Gene Expression")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the PIPseeker ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters.")
    if 'sensitivity' in de.iloc[0]['sensitivity_level']: 
        runner = rc.add_run(f" To accommodate for differences between samples and facilitate accurate cell calling, PIPseeker performs cell calling at five different sensitivity. With {de.iloc[0]['sensitivity_level']}, t")
    if 'force_' in de.iloc[0]['sensitivity_level']:
        runner = rc.add_run(f"The parameter '--force-cells' is used. T") 
    runner = rc.add_run(f"he number of cells captured ranges from {cellmin} to {cellmax} and mean reads per cell ranges from {readmin} to {readmax}. Cells with extremely low number of UMI counts were filtered out. Median genes found per cell ranges from {medgenemin} to {medgenemax} and the total number of genes detected ranges from {totgenemin} to {totgenemax}.")

    notesParagraph = doc.paragraphs[31]
    replace_workflow_image(doc)

def nopipeRunCommentsAndNotes(doc, values, de):
    rangemin = str(int(min(list(de[u'FastQC_total_sequences']))/1000000))
    rangemax = str(int(max(list(de[u'FastQC_total_sequences']))/1000000))
    rc = doc.paragraphs[17]
    rc.text = "" 
    runner = rc.add_run(f"{values['count']} Single Cell Sequencing")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. The samples have sequencing yields ranging from {rangemin} to {rangemax} million read. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run.")
    runner.bold = True 
    delete_paragraph(doc.paragraphs[20])
    replace_workflow_image(doc) ## remove the workflow image if no workflow image is provided.

def rnaFinalTable(doc, de):
    headers = ["Sample", "Estimated Number of Cells", "Mean Reads per Cell", "Median Genes per Cell", "Total Genes Detected", "Median UMI Counts per Cell"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        p.cell(i+1,0).text = str(de[u"Sample"][i])
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,1).text = "{:,}".format(int(de[u"Estimated Number of Cells"][i]))
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,2).text = "{:,}".format(int(de[u"Mean Reads per Cell"][i]))
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,3).text = "{:,}".format(int(de[u"Median Genes per Cell"][i]))
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,4).text = "{:,}".format(int(de[u"Total Genes Detected"][i]))
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,5).text = "{:,}".format(int(de[u"Median UMI Counts per Cell"][i]))
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

def pipseqFinalTable(doc, de):
    headers = ["sample", "sensitivity_level", "num_input_reads", "mapping_pct_txome", "num_cell_barcodes", "mean_reads_per_cell", "num_genes_expressed_in_cells"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        for idx, name in enumerate(headers):
            if name == "mapping_pct_txome":
                de[name][i] = '{:.2%}'.format(de[name][i])
            p.cell(i+1,idx).text = str(de[name][i])
            p.cell(i+1,idx).paragraphs[0].paragraph_format.space_before = Pt(2)
            p.cell(i+1,idx).paragraphs[0].paragraph_format.space_after = 0

def nopipeFinalTable(doc, de):
    headers = ["Sample", "FastQC_total_sequences"]
    p=doc.tables[3]
    finalTableSetup(p, headers)
    for i in range(len(de)):
        p.add_row()
        for idx, name in enumerate(headers):
            p.cell(i+1,idx).text = str(de[name][i])
            p.cell(i+1,idx).paragraphs[0].paragraph_format.space_before = Pt(2)
            p.cell(i+1,idx).paragraphs[0].paragraph_format.space_after = 0

def fbRunCommentsAndNotes(doc, values, de):
    barcodeq30 = str(float(min(list(de[u'Q30 Bases in Barcode'])))*100)
    umiq30 = str(float(min(list(de[u'Q30 Bases in UMI'])))*100)
    cellmin = "{:,}".format(int(min(list(de[u'Estimated Number of Cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'Estimated Number of Cells']))))

    if sum(['RNA Read' in i for i in de.columns ]) == 1:
        rnareadq30 = str(float(min(list(de[u'Q30 Bases in RNA Read'])))*100)
    else:
        rnareadq30 = str(round(float(min([min(de[i]) for i in de.columns if 'Q30 Bases in RNA Read' in i]))*100, 2))

    rangemin = str(int(min(list(de[u'Number of Reads']))/1000000))
    rangemax = str(int(max(list(de[u'Number of Reads']))/1000000))
    readmin = "{:,}".format(int(min(list(de[u'Mean Reads per Cell']))))
    readmax ="{:,}".format(int(max(list(de[u'Mean Reads per Cell']))))
    medgenemin = "{:,}".format(int(min(list(de[u'Median Genes per Cell']))))
    medgenemax ="{:,}".format(int(max(list(de[u'Median Genes per Cell']))))
    totgenemin = "{:,}".format(int(min(list(de[u'Total Genes Detected']))))
    totgenemax ="{:,}".format(int(max(list(de[u'Total Genes Detected']))))
    antireadmin = "{:,}".format(int(min(list(de[u'Antibody: Mean Reads per Cell']))))
    antireadmax ="{:,}".format(int(max(list(de[u'Antibody: Mean Reads per Cell']))))
    antifracreadmin = str(float(min(list(de[u'Antibody: Fraction Antibody Reads'])))*100)
    antifracreadmax = str(float(max(list(de[u'Antibody: Fraction Antibody Reads'])))*100)
    antifracreadusablemin = str(round(float(min(list(de[u'Antibody: Fraction Antibody Reads Usable'])))*100, 2))
    antifracreadusablemax = str(round(float(max(list(de[u'Antibody: Fraction Antibody Reads Usable'])))*100, 2))
    antiumimin = "{:,}".format(int(min(list(de[u'Antibody: Median UMIs per Cell (summed over all recognized antibody barcodes)']))))
    antiumimax ="{:,}".format(int(max(list(de[u'Antibody: Median UMIs per Cell (summed over all recognized antibody barcodes)']))))

    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("Gene Expression and Feature Barcode")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {rnareadq30}% of bases in the read have Q30 or above. More than {umiq30}% of bases in the UMI have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax} and mean reads per cell ranges from {readmin} to {readmax}. Cells with extremely low number of UMI counts were filtered out. Median genes found per cell ranges from {medgenemin} to {medgenemax} and the total number of genes detected ranges from {totgenemin} to {totgenemax}.\n\nThe mean antibody reads per cell ranges from {antireadmin} to {antireadmax}. The fraction of antibody library reads that contain a recognized antibody barcode range from {antifracreadmin}% to {antifracreadmax}%. The fraction of antibody reads usable ranges from {antifracreadusablemin}% to {antifracreadusablemax}%. The median antibody UMIs per cell is between {antiumimin} and {antiumimax}.")

    notesParagraph = doc.paragraphs[31]
    bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
    bulletEntry(notesParagraph, "Mean Reads per Cell", "The total number of sequenced reads divided by the estimated number of cells.")
    bulletEntry(notesParagraph, "Median Genes per Cell", "The median number of genes detected per cell barcode.")
    bulletEntry(notesParagraph, "Median UMI Counts per Cell", "The median number of UMI counts per cell barcode.")
    bulletEntry(notesParagraph, "Number of Reads", "Total number of sequenced reads.")
    bulletEntry(notesParagraph, "Reads Mapped Confidently to Transcriptome", "Fraction of reads that mapped to a unique gene in the transcriptome with a high mapping quality score as reported by the aligner.")
    bulletEntry(notesParagraph, "Valid Barcodes", "Fraction of reads with cell-­barcodes that match the whitelist.")
    bulletEntry(notesParagraph, "Sequencing Saturation", "The fraction of reads originating from an already-­observed UMI. This is a function of library complexity and sequencing depth. More specifically, this is the fraction of confidently mapped, valid cell-­barcode, valid UMI reads that had a non-­unique (cell-­barcode, UMI, gene).")

def atacRunCommentsAndNotes(doc, values, de):
    barcodeq30 = "{0:.2f}".format(float(min(list(de[u'Q30 bases in barcode'])))*100)
    cellmin = "{:,}".format(int(min(list(de[u'Estimated number of cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'Estimated number of cells']))))
    rnareadq30 = "{0:.2f}".format(min([float(min(list(de[u'Q30 bases in read 1']))), float(min(list(de[u'Q30 bases in read 2'])))])*100)
    siq30 = "{0:.2f}".format(float(min(list(de[u'Q30 bases in sample index i1'])))*100)

    rangemin = "{:,}".format(int(int(min(list(de[u'Sequenced read pairs'])))/1000000))
    rangemax = "{:,}".format(int(int(max(list(de[u'Sequenced read pairs'])))/1000000))
    numfragmentsmin = "{:,}".format(int(min(list(de[u'Median high-quality fragments per cell']))))
    numfragmentsmax = "{:,}".format(int(max(list(de[u'Median high-quality fragments per cell']))))
    minmapped = "{0:.2f}".format(float(min(list(de[u'Confidently mapped read pairs'])))*100)
    minoverlaptss = "{0:.2f}".format(float(min(list(de[u'Fraction of high-quality fragments overlapping TSS'])))*100)
    minoverlappeaks = "{0:.2f}".format(float(min(list(de[u'Fraction of high-quality fragments overlapping peaks'])))*100)


    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("ATAC")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {rnareadq30}% of bases in the read have Q30 or above, and {siq30}% or more of bases in the sample index have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax}. The median fragments per cell ranges from {numfragmentsmin} to {numfragmentsmax}. The percentage of reads mapped confidently to genome (>30 mapq) is above {minmapped}% for all the samples. The fraction of fragments (that passed all filters) overlapping TSS is above {minoverlaptss}% and overlapping peaks is above {minoverlappeaks}%.")

    notesParagraph = doc.paragraphs[31]
    bulletEntry(notesParagraph, "Number of Cell", "The total number of barcodes identified as cells.")
    bulletEntry(notesParagraph, "Sequenced Read Pairs", "Total number of sequenced read pairs assigned to the sample.")
    bulletEntry(notesParagraph, "Confidently Mapped Read Pairs", "The fraction of sequenced read pairs with mapping quality > 30.")
    bulletEntry(notesParagraph, "High-Quality Fragments", "Read pairs with a valid barcode that map to the nuclear genome with mapping quality > 30, are not chimeric and not duplicate.")
    bulletEntry(notesParagraph, "Median High-Quality Fragments per Cell", "The median number of high-quality fragments per cell barcode.")
    bulletEntry(notesParagraph, "Fraction of High-Quality Fragments Overlapping Targets", "The fraction of high-quality fragments that overlap targeted regions.")

def atacFinalTable(doc, de):
    headers = ["Sample", "Estimated Number of Cells", "Sequenced Read Pairs", "Confidently Mapped Read Pairs", "Median High-Quality Fragments per Cell", "Fraction of High-Quality Fragments Overlapping Targets"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        p.cell(i+1,0).text = str(de[u"Sample ID"][i])
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,1).text = "{:,}".format(int(de[u"Estimated number of cells"][i]))
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,2).text = "{:,}".format(int(de[u"Sequenced read pairs"][i]))
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,3).text = "{0:.4f}".format(float(de[u"Confidently mapped read pairs"][i]))
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,4).text = "{0:,}".format(float(de[u"Median high-quality fragments per cell"][i]))
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,5).text = "{0:.4f}".format(float(de[u"Fraction of high-quality fragments overlapping peaks"][i]))
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

def atacOldRunCommentsAndNotes(doc, values, de):
    barcodeq30 = "{0:.2f}".format(float(min(list(de[u'bc_q30_bases_fract'])))*100)
    #index = str(float(min(list(de[u'Q30 Bases in Sample Index'])))*100)
    cellmin = "{:,}".format(int(min(list(de[u'annotated_cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'annotated_cells']))))

    rnareadq30 = "{0:.2f}".format(min([float(min(list(de[u'r1_q30_bases_fract']))), float(min(list(de[u'r2_q30_bases_fract'])))])*100)
    siq30 = "{0:.2f}".format(float(min(list(de[u'si_q30_bases_fract'])))*100)

    version = str(list(de[u'cellranger-atac_version'])[0])

    rangemin = "{:,}".format(int(min(list(de[u'num_fragments']))))
    rangemax = "{:,}".format(int(max(list(de[u'num_fragments']))))

    numfragmentsmin = "{:,}".format(int(min(list(de[u'median_fragments_per_cell']))))
    numfragmentsmax = "{:,}".format(int(max(list(de[u'median_fragments_per_cell']))))

    usablefragmentsmin = "{:,}".format(int(min(list(de[u'total_usable_fragments']))))
    usablefragmentsmax = "{:,}".format(int(max(list(de[u'total_usable_fragments']))))

    minmapped = "{0:.2f}".format(float(min(list(de[u'frac_mapped_confidently'])))*100)
    minoverlaptargets = "{0:.2f}".format(float(min(list(de[u'frac_fragments_overlapping_targets'])))*100)


    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("ATAC")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {rnareadq30}% of bases in the read have Q30 or above, and {siq30}% or more of bases in the sample index have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax}. The median fragments per cell ranges from {numfragmentsmin} to {numfragmentsmax}. Total usable fragments per sample was between {usablefragmentsmin} and {usablefragmentsmax}. The percentage of reads mapped confidently to genome (>30 mapq) is above {minmapped}% for all the samples. The fraction of fragments (that passed all filters) overlapping targeted regions is above {minoverlaptargets}%.")

    notesParagraph = doc.paragraphs[31]
    bulletEntry(notesParagraph, "Number of Cell", "The total number of barcodes identified as cells.")
    bulletEntry(notesParagraph, "Sequenced Read Pairs", "Total number of sequenced read pairs assigned to the sample.")
    bulletEntry(notesParagraph, "Confidently Mapped Read Pairs", "The fraction of sequenced read pairs with mapping quality > 30.")
    bulletEntry(notesParagraph, "High-Quality Fragments", "Read pairs with a valid barcode that map to the nuclear genome with mapping quality > 30, are not chimeric and not duplicate.")
    bulletEntry(notesParagraph, "Median High-Quality Fragments per Cell", "The median number of high-quality fragments per cell barcode.")
    bulletEntry(notesParagraph, "Fraction of High-Quality Fragments Overlapping Targets", "The fraction of high-quality fragments that overlap targeted regions.")

def atacOldFinalTable(doc, de):
    headers = ["Sample", "Estimated Number of Cells", "Median Fragments per Cell", "Total Usable Fragments", "Fraction Mapped Confidently", "Fraction Fragments Overlapping Targets"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        p.cell(i+1,0).text = str(de[u"Sample"][i])
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,1).text = "{:,}".format(int(de[u"annotated_cells"][i]))
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,2).text = "{0:,}".format(float(de[u"median_fragments_per_cell"][i]))
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,3).text = "{:,}".format(int(de[u"total_usable_fragments"][i]))
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,4).text = "{0:f}".format(float(de[u"frac_mapped_confidently"][i]))
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,5).text = "{0:f}".format(float(de[u"frac_fragments_overlapping_targets"][i]))
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

def vdjRunCommentsAndNotes(doc, values, de):
    barcodeq30 = str(float(min(list(de[u'Q30 Bases in Barcode'])))*100)
    umiq30 = str(float(min(list(de[u'Q30 Bases in UMI'])))*100)
    cellmin = "{:,}".format(int(min(list(de[u'Estimated Number of Cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'Estimated Number of Cells']))))

    if sum(['RNA Read' in i for i in de.columns ]) == 1:
        rnareadq30 = str(float(min(list(de[u'Q30 Bases in RNA Read 1'])))*100)
    else:
        rnareadq30 = str(round(float(min([min(de[i]) for i in de.columns if 'Q30 Bases in RNA Read' in i]))*100, 2))

    rangemin = "{:,}".format(int(min(list(de[u'Number of Read Pairs']))/1000000))
    rangemax = "{:,}".format(int(max(list(de[u'Number of Read Pairs']))/1000000))
    readmin = "{:,}".format(int(min(list(de[u'Mean Read Pairs per Cell']))))
    readmax = "{:,}".format(int(max(list(de[u'Mean Read Pairs per Cell']))))
    vjspanmin = str(float(min(list(de[u'Cells With Productive V-J Spanning Pair'])))*100)
    vjspanmax = str(float(max(list(de[u'Cells With Productive V-J Spanning Pair'])))*100)
    clonomin = "{:,}".format(float(min(list(de[u'Paired Clonotype Diversity']))))
    clonomax = "{:,}".format(float(max(list(de[u'Paired Clonotype Diversity']))))

    rc = doc.paragraphs[17]
    rc.text = ""
    runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run("VDJ")
    runner.bold = True
    runner = rc.add_run(" libraries were sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {rnareadq30}% of bases in RNA reads have Q30 or above. More than {umiq30}% of bases in the UMI have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax} and mean reads per cell ranges from {readmin} to {readmax}. Cells with extremely low number of UMI counts were filtered out. Cells with productive V-J spanning pair ranges from {vjspanmin}% to {vjspanmax}%. Paired clonotype diversity is between {clonomin} and {clonomax}.")

    #bulletEntry(doc, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
    #bulletEntry(doc, "Mean Reads per Cell", "The total number of sequenced reads divided by the estimated number of cells.")
    #bulletEntry(doc, "Number of Reads", "Total number of sequenced reads.")
    #bulletEntry(doc, "Number of Cells with Productive V-J Spanning Pair", "Fraction of cell-associated barcodes with at least one productive contig for each chain of the receptor pair.")
    #bulletEntry(doc, "Reads Mapped to Any V(D)J Gene", "Fraction of reads that partially or wholly map to any germline V(D)J gene segment.")
    #bulletEntry(doc, "Paired Clonotype Diversity", "Effective diversity of the paired clonotypes, computed as the Inverse Simpson Index of the clonotype frequencies.")

    notesParagraph = doc.paragraphs[31]
    bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
    bulletEntry(notesParagraph, "Mean Reads per Cell", "The total number of sequenced reads divided by the estimated number of cells.")
    bulletEntry(notesParagraph, "Number of Reads", "Total number of sequenced reads.")
    bulletEntry(notesParagraph, "Number of Cells with Productive V-J Spanning Pair", "Fraction of cell-associated barcodes with at least one productive contig for each chain of the receptor pair.")
    bulletEntry(notesParagraph, "Reads Mapped to Any V(D)J Gene", "Fraction of reads that partially or wholly map to any germline V(D)J gene segment.")
    bulletEntry(notesParagraph, "Paired Clonotype Diversity", "Effective diversity of the paired clonotypes, computed as the Inverse Simpson Index of the clonotype frequencies.")

def vdjFinalTable(doc, de):
    headers = ["Sample", "Estimated Number of Cells", "Mean Reads per Cell", "Number of Cells with Productive V-J Spanning Pair", "Reads Mapped to Any V(D)J Gene", "Paired Clonotype Diversity"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        p.cell(i+1,0).text = str(de[u"Sample"][i])
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,1).text = "{:,}".format(int(de[u"Estimated Number of Cells"][i]))
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,2).text = "{:,}".format(int(de[u"Mean Read Pairs per Cell"][i]))
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,3).text = "{:,}".format(int(de[u"Number of Cells With Productive V-J Spanning Pair"][i]))
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,4).text = str(float(de[u"Reads Mapped to Any V(D)J Gene"][i])*100) + "%"
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,5).text = "{:,}".format(float(de[u"Paired Clonotype Diversity"][i]))
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

def multiRunCommentsAndNotes(doc, values, data, runInfo, libraries):
    de = {i:data.parse(i)for i in data.sheet_names}
    de_sample = data.parse('Sample')

    rc = doc.paragraphs[17]
    rc.text = ""

    prev = False
    if any('Gene Expression' in i for i in de['Library']):
        rangemin = str(int(min(list(de['Library'][u'Gene Expression Number of reads']))/1000000))
        cellmin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Cells']))))
        cellmax = "{:,}".format(int(max(list(de['Sample'][u'Gene Expression Cells']))))
        readmin = "{:,}".format(int(min(list(de['Library'][u'Gene Expression Mean reads per cell']))))
        readmax ="{:,}".format(int(max(list(de['Library'][u'Gene Expression Mean reads per cell']))))
        medgenemin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Median genes per cell']))))
        medgenemax ="{:,}".format(int(max(list(de['Sample'][u'Gene Expression Median genes per cell']))))
        totgenemin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Total genes detected']))))
        totgenemax ="{:,}".format(int(max(list(de['Sample'][u'Gene Expression Total genes detected']))))

        runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
        runner = rc.add_run("Gene Expression")
        runner.bold = True
        if libraries == None:
            runner = rc.add_run(" libraries were sequenced on a ")
            runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
            runner.bold = True
            runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
            runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
            runner.bold = True
        else:
            with open(libraries, 'r') as f:
                flowcells = set([line.strip().split(',')[1] for line in f if 'Gene Expression' in line.strip().split(',')[-1]])
                runner = rc.add_run(" libraries were sequenced on ")
                platform = dict()
                platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                uniq = list(set(platforms))
                uniq.sort()
                runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                runner.bold = True
                runner = rc.add_run(f". All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as ")
                runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                runner.bold = True
        runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
        runner = rc.add_run(f"{values['version']}")
        runner.bold = True
        runner = rc.add_run(f" software using the default parameters. For each sample the number of cells captured ranges from {cellmin} to {cellmax} and mean reads per cell ranges from {readmin} to {readmax}. Cells with extremely low number of UMI counts were filtered out. Median genes found per cell ranges from {medgenemin} to {medgenemax} and the total number of genes detected ranges from {totgenemin} to {totgenemax}.")

        notesParagraph = doc.paragraphs[31]
        bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
        bulletEntry(notesParagraph, "Median Reads per Cell", "Median number of read pairs sequenced from the cells assigned to this sample.")
        bulletEntry(notesParagraph, "Median Genes per Cell", "The median number of genes detected per cell assigned to this sample. Detection is defined as the presence of at least 1 UMI count.")
        bulletEntry(notesParagraph, "Median UMI Counts per Cell", "The median number of UMI counts per cell barcode.")
        bulletEntry(notesParagraph, "Reads Mapped Confidently to Transcriptome", "Fraction of reads that mapped to a unique gene in the transcriptome with a high mapping quality score as reported by the aligner.")
        bulletEntry(notesParagraph, "Sequencing Saturation", "The fraction of reads originating from an already-­observed UMI. This is a function of library complexity and sequencing depth. More specifically, this is the fraction of confidently mapped, valid cell-­barcode, valid UMI reads that had a non-­unique (cell-­barcode, UMI, gene).")

        prev = True

    if any('VDJ' in i for i in de['Library']):
        if prev:
            rc.add_run("\n\n")
        fil_de_lib = de['Library'].loc[:, de['Library'].columns.str.contains('VDJ')].dropna()
        fil_de_sam = de['Sample'].loc[:, de['Sample'].columns.str.contains('VDJ')].dropna()
        rangemin = str(int(min([min(list(fil_de_lib[i])) for i in fil_de_lib if 'Number of reads' in i])/1000000))
        cellmin = "{:,}".format(int(min([min(list(fil_de_lib[i])) for i in fil_de_lib if 'Estimated number of cells' in i])))
        cellmax = "{:,}".format(int(max([max(list(fil_de_lib[i])) for i in fil_de_lib if 'Estimated number of cells' in i])))
        readmin = "{:,}".format(int(min([min(list(fil_de_lib[i])) for i in fil_de_lib if 'Mean reads per cell' in i])))
        readmax = "{:,}".format(int(max([max(list(fil_de_lib[i])) for i in fil_de_lib if 'Mean reads per cell' in i])))
        vjspanmin = "{:,}".format(int(min([min(list(fil_de_sam[i])) for i in fil_de_sam if 'Number of cells with productive V-J spanning pair' in i])))
        vjspanmax = "{:,}".format(int(max([max(list(fil_de_sam[i])) for i in fil_de_sam if 'Number of cells with productive V-J spanning pair' in i])))
        clonomin = "{:,}".format(float(min([min(list(fil_de_sam[i])) for i in fil_de_sam if 'Paired clonotype diversity' in i])))
        clonomax = "{:,}".format(float(max([max(list(fil_de_sam[i])) for i in fil_de_sam if 'Paired clonotype diversity' in i])))
        runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
        runner = rc.add_run("VDJ")
        runner.bold = True
        if libraries == None:
            runner = rc.add_run(" libraries were sequenced on a ")
            runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
            runner.bold = True
            runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
            runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
            runner.bold = True
        else:
            with open(libraries, 'r') as f:
                flowcells = set([line.strip().split(',')[1] for line in f if 'VDJ' in line.strip().split(',')[-1]])
                runner = rc.add_run(" libraries were sequenced on ")
                platform = dict()
                platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                uniq = list(set(platforms))
                uniq.sort()
                runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                runner.bold = True
                runner = rc.add_run(f". All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as ")
                runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                runner.bold = True
        runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
        runner = rc.add_run(f"{values['version']}")
        runner.bold = True
        runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax} and mean reads per cell ranges from {readmin} to {readmax}. Cells with productive V-J spanning pair ranges from {vjspanmin} to {vjspanmax}. Paired clonotype diversity is between {clonomin} and {clonomax}.")

        notesParagraph = doc.paragraphs[31]
        bulletEntry(notesParagraph, "Number of Cells with Productive V-J Spanning Pair", "Fraction of cell-associated barcodes with at least one productive contig for each chain of the receptor pair.")
        bulletEntry(notesParagraph, "Reads Mapped to Any V(D)J Gene", "Fraction of reads that partially or wholly map to any germline V(D)J gene segment.")
        bulletEntry(notesParagraph, "Paired Clonotype Diversity", "Effective diversity of the paired clonotypes, computed as the Inverse Simpson Index of the clonotype frequencies.")
        prev = True

    if any('Antibody Capture' in i for i in de['Library']):
        if prev:
            rc.add_run("\n\n")
            rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Number of reads' in i])/1000000))
            cellmin = "{:,}".format(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Estimated number of cells' in i])))
            cellmax = "{:,}".format(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Estimated number of cells' in i])))
            antireadmin = "{:,}".format(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Mean reads per cell' in i])))
            antireadmax = "{:,}".format(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Mean reads per cell' in i])))
            antifracreadmin = "{:,}".format(float(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Fraction antibody reads usable' in i]))*100)
            antifracreadmax = "{:,}".format(float(max([max(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Fraction antibody reads usable' in i]))*100)
            antifracreadusablemin = "{:,}".format(int(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Mean antibody reads usable per cell' in i])))
            antifracreadusablemax = "{:,}".format(int(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Mean antibody reads usable per cell' in i])))
            antiumimin = "{:,}".format(float(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
            antiumimax = "{:,}".format(float(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
            runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
            runner = rc.add_run("Antibody")
            runner.bold = True
            if libraries == None:
                runner = rc.add_run(" libraries were sequenced on a ")
                runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
                runner.bold = True
                runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
                runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
                runner.bold = True
            else:
                with open(libraries, 'r') as f:
                    flowcells = set([line.strip().split(',')[1] for line in f if 'Antibody' in line.strip().split(',')[-1]])
                    runner = rc.add_run(" libraries were sequenced on ")
                    platform = dict()
                    platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                    uniq = list(set(platforms))
                    uniq.sort()
                    runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                    runner.bold = True
                    runner = rc.add_run(f". All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as ")
                    runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                    runner.bold = True
            runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
            runner = rc.add_run(f"{values['version']}")
            runner.bold = True
            runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax}. The mean antibody reads per cell ranges from {antireadmin} to {antireadmax}. The fraction of antibody library reads that contain a recognized antibody barcode range from {antifracreadmin}% to {antifracreadmax}%. The number of antibody reads usable per cell ranges from {antifracreadusablemin} to {antifracreadusablemax}. The median antibody UMIs per cell is between {antiumimin} and {antiumimax}.")
            prev = True

    if any('CRISPR Guide Capture' in i for i in de['Library']):
        if prev:
            rc.add_run("\n\n")
            rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'CRISPR Guide Capture' in i and 'Number of reads' in i])/1000000))
            rangemax = str(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'CRISPR Guide Capture' in i and 'Number of reads' in i])/1000000))
            meanpercellmin = "{:,}".format(int(min([min([j for j in list(de['Library'][i])]) for i in de['Library'] if 'CRISPR Guide Capture' in i and 'Mean reads per cell' in i])))
            meanpercellmax = "{:,}".format(int(max([max([j for j in list(de['Library'][i])]) for i in de['Library'] if 'CRISPR Guide Capture' in i and 'Mean reads per cell' in i])))
            pctprotospacermin = "{:.2%}".format(float(min([min([j for j in list(de['Sample'][i])]) for i in de['Sample'] if 'CRISPR Guide Capture' in i and 'with one or more protospacers detected' in i])))
            pctprotospacermax = "{:.2%}".format(float(max([max([j for j in list(de['Sample'][i])]) for i in de['Sample'] if 'CRISPR Guide Capture' in i and 'with one or more protospacers detected' in i])))
            pct2protospacermin = "{:.2%}".format(float(min([min([j for j in list(de['Sample'][i])]) for i in de['Sample'] if 'CRISPR Guide Capture' in i and 'with two or more protospacers detected' in i])))
            pct2protospacermax = "{:.2%}".format(float(max([max([j for j in list(de['Sample'][i])]) for i in de['Sample'] if 'CRISPR Guide Capture' in i and 'with two or more protospacers detected' in i])))
            runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
            runner = rc.add_run("CRISPR Guide Capture")
            runner.bold = True
            runner = rc.add_run(" libraries were sequenced on a ")
            runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
            runner.bold = True
            runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
            runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run. ")
            runner.bold = True
            print()
            runner = rc.add_run(f"The mean numbers of reads per cell range from {meanpercellmin} to {meanpercellmax}. The percentages of cells with one or more protospacers detected range from {pctprotospacermin} to {pctprotospacermax}. The percentages of cells with two or more protospacers detected range from {pct2protospacermin} to {pct2protospacermax}.")
            prev = True
        
        if any('Multiplexing Capture' in i for i in de['Library']):
            if prev:
                rc.add_run("\n\n")
                rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Number of reads' in i])/1000000))
                cellassignmin = "{:,}".format(int(min([min([j.split(' ')[0] for j in list(de['Library'][i])]) for i in de['Library'] if 'Multiplexing Capture' in i and 'Cells assigned to a sample' in i])))
                cellassignmax = "{:,}".format(int(max([max([j.split(' ')[0] for j in list(de['Library'][i])]) for i in de['Library'] if 'Multiplexing Capture' in i and 'Cells assigned to a sample' in i])))
                sampleassignmin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Samples assigned at least one cell' in i])))
                fracreadsmin = "{:}".format(float(min([min(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Fraction reads in cell-associated barcodes' in i]))*100)
                fracreadsmax = "{:,}".format(float(max([max(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Fraction reads in cell-associated barcodes' in i]))*100)
                antifracreadusablemin = "{:,}".format(int(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Antibody reads usable per cell' in i])))
                antifracreadusablemax = "{:,}".format(int(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Antibody reads usable per cell' in i])))
                antiumimin = "{:,}".format(float(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
                antiumimax = "{:,}".format(float(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
                runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
                runner = rc.add_run("Multiplexing")
                runner.bold = True
                if libraries == None:
                    runner = rc.add_run(" libraries were sequenced on a ")
                    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
                    runner.bold = True
                    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
                    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
                    runner.bold = True
                else:
                    with open(libraries, 'r') as f:
                        flowcells = set([line.strip().split(',')[1] for line in f if 'Multiplexing' in line.strip().split(',')[-1]])
                        runner = rc.add_run(" libraries were sequenced on ")
                        platform = dict()
                        platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                        uniq = list(set(platforms))
                        uniq.sort()
                        runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                        runner.bold = True
                        runner = rc.add_run(f". All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as ")
                        runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                        runner.bold = True
                runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
                runner = rc.add_run(f"{values['version']}")
                runner.bold = True
                runner = rc.add_run(f" software using the default parameters. The number of cells assigned to a sample ranges from {cellassignmin} to {cellassignmax}. There were a minimum of {sampleassignmin} samples assigned at least one cell. The fraction of reads in cell-associated barcodes range from {fracreadsmin}% to {fracreadsmax}%.")
                prev = True

def multiFinalTable(doc, data):
    de = {i:data.parse(i)for i in data.sheet_names}
    de = de['Sample']
    de_sample = data.parse('Sample')
    headers = ["Sample", "Gene Expression Cells", "Gene Expression Median reads per cell", "Gene Expression Median genes per cell", "Gene Expression Total genes detected", "Gene Expression Median UMI counts per cell"]
    p=doc.tables[3]
    finalTableSetup(p, headers)

    for i in range(len(de)):
        p.add_row()
        p.cell(i+1,0).text = str(de[u"Sample"][i])
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,1).text = "{:,}".format(int(de[u"Gene Expression Cells"][i]))
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,2).text = "{:,}".format(int(de[u"Gene Expression Median reads per cell"][i]))
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,3).text = "{:,}".format(int(de[u"Gene Expression Median genes per cell"][i]))
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,4).text = "{:,}".format(int(de[u"Gene Expression Total genes detected"][i]))
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p.cell(i+1,5).text = "{:,}".format(int(de[u"Gene Expression Median UMI counts per cell"][i]))
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
        p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
        p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT

parser = argparse.ArgumentParser(description="""Generate a single cell word report based on the library type and flowcell information""")
parser.add_argument("-e", "--excel", metavar="ExcelReport", dest="excelfile",
                    action="store", type=str, required=True,
                    help="Summary Excel file containing the CellRanger summary")
parser.add_argument("-m", "--metadata", metavar="Metadata.txt", dest="metafile",
                    action="store", type=str, required=True,
                    help="Metadata file associated with the project")
parser.add_argument("-c", "--cellranger", metavar="cellranger-6.0.2", dest="cellranger",
                    action="store", type=str, required=True,
                    help="Cell Ranger version used with version number")
parser.add_argument("-p", "--pipeline", metavar="rna", dest="pipeline",
                    action="store", type=str.lower, required=True,
                    help="Analysis version used: rna, vdj, multiome, fb")
parser.add_argument("-r", "--run_names", metavar="210728_A00430_0443_BHFH52DRXY", dest="run_names",
                    action="store", type=str, required=False,
                    help="List of run names, comma separated if more than one is added")
parser.add_argument("-l", "--libraries", metavar="libraries.csv", dest="libraries",
                    action="store", type=str, required=False,
                    help="Libraries file used to set up the run")
parser.add_argument("-y", "--yields", dest="yields", type=float, 
                    help="Total yields of sequencing data")
parser.add_argument("-t", "--tracking-sheet", dest="tracking_sheet", 
                    default = "/mnt/ccrsf-ifx/Report_archive/processed_ccrsfifx_SingleCell.csv",
                    help="Output for trakcing information")
parser.add_argument("--test", dest="test", action = 'store_true', 
                    help="Allow script to run in test mode and output tracking info into test csv file")
parser.add_argument("-w", "--workflow", type=str, required=False,
                    help="Workflow image to replace the one in the word")
args = parser.parse_args()

if args.pipeline == "pipseq": 
    if args.workflow == None: 
        sys.stderr.write("\nWorkflow image (-w) is required for PIPseq data \n\n")
        sys.exit() 
    else:
        if not os.path.exists(args.workflow):
            sys.stderr.write("\nWorkflow image not detected at " + args.workflow + ". Please check it out! \n\n")
            sys.exit()


# excelfile = sys.argv[1]
# metafile = sys.argv[2]
# cellranger = sys.argv[3]
# pipeline = sys.argv[4] #Multiome
# run_names = sys.argv[5].split(',')
excelfile = args.excelfile
metafile = args.metafile
cellranger = args.cellranger
pipeline = args.pipeline
if args.run_names == None:
    run_names = None
else:
    run_names = args.run_names.split(',')

filenm=os.path.splitext(metafile)[0][:-9]+'.docx'
emailfile = os.path.splitext(metafile)[0][:-9]+'_email.docx'
data = pd.ExcelFile(excelfile)

metadata = pd.read_csv(metafile, sep='\t')


if str(data.sheet_names[0]) != "metrics_summary":
    de = data.parse(data.sheet_names[0])
else:
    de = data.parse('metrics_summary')


docfile = os.path.join(path2script, "data/10X_sc.docx") 
print(docfile)
doc = Document(docfile)

values = dict()
values = getCellRangerVersion(values, de, cellranger, pipeline)
values['count'] = str(len(de))
values['pi'] = metadata[metadata.columns[1]][0]
values['lc'] = metadata[metadata.columns[2]][0]
values['bc'] = metadata[metadata.columns[4]][0]
values['pt'] = metadata[metadata.columns[3]][0]
values['csac'] = str(metadata[metadata.columns[10]][0])
values['stip'] = str(len(metadata))
values['sitr'] = str(len(de))
values['date'] = time.strftime("%x")

replace_string(doc,"{PI}", values['pi'])
replace_string(doc,"{LC}", values['lc'])
replace_string(doc,"{BC}", values['bc'])
replace_string(doc,"{PT}", values['pt'])
replace_string(doc,"{CSAC}", values['csac'])
replace_string(doc,"{STIP}", values['stip'])
replace_string(doc,"{SITR}", values['sitr'])
replace_string(doc,"{DATE}", values['date'])
dme_link = generate_DME_link(values['pt'])
replace_string_hyperlink(doc,"{DMELINK}", dme_link)

values['seqType'] = metadata['Application'][0]

#st = str(len(summary)) + ' ' + str(metadata[metadata.columns[8]][0])
values['platform'] = str(metadata[metadata.columns[11]][0])

runInfo = dict()
if run_names != None:
    runInfo = getRunInfo(run_names)

if len(runInfo.keys()) != 0:
    flowcell = list(runInfo.keys())[0]
    for i in runInfo[flowcell]:
        values[i] = runInfo[flowcell][i]
    if runInfo[flowcell]['read1'] == runInfo[flowcell]['read2']:
        values['sym'] = 'symmetric'
    else:
        values['sym'] = 'non-symmetric'
else:
    if 'x' in str(metadata['ReadLength'][0]):
        if '2x' in metadata['ReadLength']:
            values['sym'] = 'symmetric'
            values['read1'] = metadata['ReadLength'][0].split('x')[1]
            values['read2'] = read1
        else:
            values['sym'] = 'non-symmetric'
            values['read1'] = metadata['ReadLength'][0].split('x')[1]
            values['read2'] = '0'
    else:
    	temp = re.findall(": *(\d+)", metadata['ReadLength'][0])
    	values['read1'] = temp[0]
    	values['read2'] = temp[1]
    	if values['read1'] == values['read2']:
    	    values['sym'] = 'symmetric'
    	else:
    	    values['sym'] = 'non-symmetric'
    values['flowcell'] = str(metafile.split('_')[-2])

table1 = doc.tables[0]
if len(runInfo) > 0:
    table1.cell(0,2).text = ', '.join(runInfo[i]['flowcell'] for i in runInfo)
    table1.cell(1,2).text = ', '.join(set([runInfo[i]['platform'] for i in runInfo]))

    flowcellType = list()
    for i in runInfo:
        if 'xp' in runInfo[i].get('workflow', '').lower():
            workflow = "XP"
        else:
            workflow = ''
        flowcellType.append(' '.join([runInfo[i]['chemistry'], workflow]).strip())
    flowcellType.sort()
    flowcellUnique = list(set(flowcellType))
    flowcellUnique.sort()
    table1.cell(2,2).text = ', '.join([f"{flowcellType.count(i)} {i}" for i in flowcellUnique])

    runLengths = []
    for i in runInfo:
        if 'idx2' == 0:
            runLengths.append(f"R1: {runInfo[i]['read1']}, i7: {runInfo[i]['idx1']}, R2: {runInfo[i]['read2']}")
        else:
            runLengths.append(f"R1: {runInfo[i]['read1']}, i7: {runInfo[i]['idx1']}, i5: {runInfo[i]['idx2']}, R2: {runInfo[i]['read2']}")
    if len(runInfo) > 1:
        table1.cell(4,2).text = '\n'.join([f"{runInfo[i]['flowcell']}: R1: {runInfo[i]['read1']}, i7: {runInfo[i]['idx1']}, i5: {runInfo[i]['idx2']}, R2: {runInfo[i]['read2']}" for i in runInfo])
    else:
        table1.cell(4,2).text = f"R1: {runInfo[i]['read1']}, i7: {runInfo[i]['idx1']}, i5: {runInfo[i]['idx2']}, R2: {runInfo[i]['read2']}"
else:
    table1.cell(0,2).text = values['flowcell']
    table1.cell(1,2).text = str(metadata[metadata.columns[11]][0])
    table1.cell(2,2).text = str(metadata[metadata.columns[11]][0])
    table1.cell(4,2).text = str(metadata[metadata.columns[12]][0])
table1.cell(0,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(0,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(1,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(2,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(2,2).paragraphs[0].paragraph_format.space_after = 0
#table1.cell(2,2).text = str(metadata[metadata.columns[8]][0])
table1.cell(3,2).text = str(values['seqType'])
table1.cell(3,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(3,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(3,2).paragraphs[0].runs[0].font.bold = True
table1.cell(4,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(4,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(4,2).paragraphs[0].runs[0].font.bold = True
#table1.cell(4,5).text = metadata[metadata.columns[9]][0]

table1.cell(2,5).text = getLibraryProtocol(pipeline)
table1.cell(2,5).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(2,5).paragraphs[0].paragraph_format.space_after = 0
table1.cell(2,5).paragraphs[0].runs[0].font.bold = True

table1.cell(3,5).text = getReference(metadata['ReferenceGenome'][0], pipeline)
table1.cell(3,5).paragraphs[0].runs[0].font.bold = True
table1.cell(3,5).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(3,5).paragraphs[0].paragraph_format.space_after = 0

p = doc.tables[2]
if len(runInfo) > 0:
    p.cell(1,1).text = p.cell(1,1).text.replace("{RTAVERSION}", ', '.join(set([runInfo[i]['rta'] for i in runInfo])))
else:
    p.cell(1,1).text = p.cell(1,1).text.replace("{RTAVERSION}", '3.4.4')

p.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
p.cell(1,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p.cell(3,1).text = p.cell(3,1).text.replace("{VERSION}", values['version'])
if pipeline.lower() == 'pipseq':
    p.cell(3,1).text = p.cell(3,1).text.replace("Cellranger", 'PIPseeker')
else:
    p.cell(2,1).text = "cellranger mkfastq\n(Bcl2fastq 2.20)"
p.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
p.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p.cell(3,2).text = getCellRangerCommand(pipeline)
p.cell(3,2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

if pipeline.lower() == 'multiome':
    multiomeFinalTable(doc, de)
    multiomeRunCommentsAndNotes(doc, values, de)
elif pipeline.lower() == 'rna':
    rnaFinalTable(doc, de)
    ## only one sample is included in the project. 
    if de.shape[0] == 1:
        rnaRunCommentsAndNotes1Sample(doc, values, de)
    ## more than one sample included in the project
    elif de.shape[0] > 1:
        rnaRunCommentsAndNotes(doc, values, de)
    else:  # no sample detected in the excel file
        sys.stderr.write("No sample detected in " + args.excelfile + ". Please check it out! ")
        sys.exit()
elif pipeline.lower() == 'pipseq':
    pipseqFinalTable(doc, de)
    # use sensitivity_3 level to generate the run comment
    if 'sensitivity' in de.iloc[0]['sensitivity_level']:
        de_sensitivity_3 = de.query("sensitivity_level == 'sensitivity_3'")
    #pipseqRunCommentsAndNotes1Sample(doc, values, de_sensitivity_3)
    pipseqRunCommentsAndNotes(doc, values, de)
elif pipeline.lower() == 'fb':
    rnaFinalTable(doc, de)
    fbRunCommentsAndNotes(doc, values, de)
elif pipeline.lower() == 'vdj':
    vdjFinalTable(doc, de)
    vdjRunCommentsAndNotes(doc, values, de)
elif pipeline.lower() == 'atac':
    atacFinalTable(doc, de)
    atacRunCommentsAndNotes(doc, values, de)
elif pipeline.lower() == 'atac_1.2.0':
    atacOldFinalTable(doc, de)
    atacOldRunCommentsAndNotes(doc, values, de)
elif pipeline.lower() == 'multi':
    de_sample = data.parse('Sample')
    ## only one sample is included in the project. 
    if de_sample.shape[0] == 1: 
        sys.stderr.write("Only one sample detected in " + args.excelfile + "\n") 
        multiRunCommentsAndNotes1Sample(doc, values, data, runInfo, args.libraries)
    ## more than one sample included in the project
    elif de_sample.shape[0] > 1:
        multiRunCommentsAndNotes(doc, values, data, runInfo, args.libraries)
    else:  # no sample detected in the excel file
        sys.stderr.write("No sample detected in " + args.excelfile + ". Please check it out! ")
        sys.exit()
    multiFinalTable(doc, data)
elif pipeline.lower() == 'nopipe':
    nopipeFinalTable(doc, de)
    nopipeRunCommentsAndNotes(doc, values, de)
    # remove the last row (cell ranger info) of table 2
    table = doc.tables[2]
    table._tbl.remove(table.rows[-1]._tr)

doc.save(filenm)

docfile = '/mnt/ccrsf-ifx/Software/scripts/lib/wordreporttemp/ccremail.docx'
doc = Document(docfile)
app = str(metadata['Application'][0])
#pilast = values['pi'].replace("_", " ").split()[1]
if '_' in str(values['pi']):
    pilast = values['pi'].replace("_", " ").split()[1]
elif ',' in str(values['pi']):
    pilast = values['pi'].split(', ')[0]
else:
    pilast = values['pi']
st = str(len(de)) + ' ' + str(values['seqType'])
#ttext = doc.paragraphs[18].text

replace_string(doc,"{App}", app)
replace_string(doc,"{Flowcell1}", values['flowcell'])
replace_string(doc,"Pilast", pilast)
replace_string(doc,"{ST}", st)
replace_string(doc,"{CSAC}", values['csac'])
replace_string(doc,"{Flowcell2}", values['flowcell'])
replace_string(doc,"{MT}", values['platform'])
replace_string(doc,"{DMELINK}", dme_link)

#doc.paragraphs[14].text = ttext
rows,columns = de.shape
tb=doc.tables[0].rows[0].cells[0].add_table(1, columns)
tb.style = 'Table Grid'

h = list(de)
head = tb.rows[0].cells
for idx, name in enumerate(h):
    paragraph = head[idx].paragraphs[0]
    run = paragraph.add_run(name)
    run.bold = True

for r in range(rows):
    cells = tb.add_row().cells
    for col in range(columns):
        if str(de.iat[r,col]).count('.') == 1:
            try:
                #print("{0:.3f}".format(de.iat[r,col]))
                cells[col].text = "{0:.3f}".format(de.iat[r,col])
            except:
                cells[col].text = str(de.iat[r,col])
        else:
            cells[col].text = str(de.iat[r,col])

doc.save(emailfile)

def mailto(message):
    msg = MIMEText(message)
    msg['Subject'] = "Tracking sheet for single cell sequencing data analysis"
    me = "CCRSF_IFX@nih.gov"
    msg['From'] = me
    you = ['ccrsf_ifx@nih.gov', 'shaojun.xie@nih.gov']
    msg['To'] = "ccrsf_ifx@nih.gov"
    s = smtplib.SMTP('localhost')
    s.sendmail(me, you, msg.as_string())
    s.quit()
    return 1

def tracking_info4singlecell():
    csvrecord = args.tracking_sheet
    if args.test == True:
        csvrecord = csvrecord[:-4] + ".test" + csvrecord[-4:]
    out = ""
    out += re.sub(",", ";", args.run_names) + ","
    proj_name = values['pt'] # os.getcwd().split("/")[-1]
    out += proj_name + ","
    pattern = r"CS\d{6}"  
    #if not bool(re.search(pattern, proj_name)) and args.test == False:
    #    mailto(f"Project name ({proj_name}) doesn't have a valid NAS request ID. Please double check and make sure the analysis folder name has NAS ID (e.g. /path/JasonBrenchley_CS123456_9scRNASeq_020723\n{out})\n")
    #    return 1 
    #out += os.path.basename(excelfile).split('.')[0] + ","
    adate = time.strftime('%m/%d/%y', time.gmtime(os.path.getctime(excelfile)))
    out += str(adate) + ",cronjob,," + str(adate) + "," + str(args.yields/1000) + ","
    sample_num = de.iloc[:, 0].nunique() 
    if args.pipeline == "multi":
        sample_num = de.iloc[:, 0].nunique()* 2
    out += str(sample_num) + ","
    out += str(metadata['Application'][0]) + ",,"
    out += str(adate) + "\n"

    ck = open (csvrecord, 'r')
    ckk = ck.read()
    ck.close()

    if args.pipeline == "multi" and args.test == False:
        mailto(f"""
Pipeline 'multi' is used. Sample number reported in the tracking sheet might not correct:\n{out}\n
If the number above is incorrect, please mannually update the tracking information in the file below:\n
{args.tracking_sheet}
""")
    if ckk.find(proj_name) > 0:
        if ckk.find(out) == -1 and args.test == False:
            mailto(f"""
Discrepancy is found in the tracking sheet: \n{out}\n
Please mannually check the tracking information in the file below:\n
{args.tracking_sheet}
""")
    else:
        cr = open (csvrecord, 'a')
        cr.write(out)
        cr.close()

tracking_info4singlecell()
