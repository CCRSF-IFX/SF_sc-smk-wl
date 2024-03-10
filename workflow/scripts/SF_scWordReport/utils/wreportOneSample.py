# xies4 on 2022/09/06: Add multiRunCommentsAndNotes1Sample for one sample word report

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def bulletEntry(rc, header, description):
    paragraph = rc.insert_paragraph_before(style='List Bullet')
    runner = paragraph.add_run(header)
    runner.bold = True
    runner = paragraph.add_run(f" - {description}")
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.left_indent = Inches(0.25)


def rnaRunCommentsAndNotes1Sample(doc, values, de):
    barcodeq30 = str(float(min(list(de[u'Q30 Bases in Barcode'])))*100)
    umiq30 = str(float(min(list(de[u'Q30 Bases in UMI'])))*100)
    cellmin = "{:,}".format(int(min(list(de[u'Estimated Number of Cells']))))
    cellmax = "{:,}".format(int(max(list(de[u'Estimated Number of Cells']))))

    if sum(['RNA Read' in i for i in de.columns ]) == 1:
        rnareadq30 = str(float(min(list(de[u'Q30 Bases in RNA Read'])))*100)
    else:
        rnareadq30 = str(round(float(min([min(de[i]) for i in de.columns if 'Q30 Bases in RNA Read' in i]))*100, 2))

    rangemin = str(int(min(list(de[u'Number of Reads']))/1000000))
    rangemax = str(int(max(list(de[u'Number of Reads']))/1000000))
    readmin = "{:,}".format(int(min(list(de[u'Mean Reads per Cell']))))
    readmax ="{:,}".format(int(max(list(de[u'Mean Reads per Cell']))))
    medgenemin = "{:,}".format(int(min(list(de[u'Median Genes per Cell']))))
    medgenemax ="{:,}".format(int(max(list(de[u'Median Genes per Cell']))))
    totgenemin = "{:,}".format(int(min(list(de[u'Total Genes Detected']))))
    totgenemax ="{:,}".format(int(max(list(de[u'Total Genes Detected']))))

    rc = doc.paragraphs[16]
    rc.text = ""
    #runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
    runner = rc.add_run(f"One 10x Genomics Single Cell ")
    runner = rc.add_run("Gene Expression")
    runner.bold = True
    runner = rc.add_run(" library was sequenced on a ")
    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
    runner.bold = True
    runner = rc.add_run(f" run. The sample has a sequencing yield of {rangemin} million read. The sequencing run was setup as a ")
    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
    runner.bold = True
    runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. Sequencing quality is good, over {barcodeq30}% of bases in the barcode regions have Q30 or above and at least {rnareadq30}% of bases in the read have Q30 or above. More than {umiq30}% of bases in the UMI have Q30 or above. The analysis was performed with the Cell Ranger ")
    runner = rc.add_run(f"{values['version']}")
    runner.bold = True
    runner = rc.add_run(f" software using the default parameters. The number of cells captured is {cellmin} and the mean number of reads per cell is {readmin}. Cells with extremely low number of UMI counts were filtered out. The median number of genes found per cell is {medgenemin} and the total number of genes detected is {totgenemin}.")

    notesParagraph = doc.paragraphs[30]
    bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
    bulletEntry(notesParagraph, "Mean Reads per Cell", "The total number of sequenced reads divided by the estimated number of cells.")
    bulletEntry(notesParagraph, "Median Genes per Cell", "The median number of genes detected per cell barcode.")
    bulletEntry(notesParagraph, "Median UMI Counts per Cell", "The median number of UMI counts per cell barcode.")
    bulletEntry(notesParagraph, "Number of Reads", "Total number of sequenced reads.")
    bulletEntry(notesParagraph, "Reads Mapped Confidently to Transcriptome", "Fraction of reads that mapped to a unique gene in the transcriptome with a high mapping quality score as reported by the aligner.")
    bulletEntry(notesParagraph, "Valid Barcodes", "Fraction of reads with cell-­barcodes that match the whitelist.")
    bulletEntry(notesParagraph, "Sequencing Saturation", "The fraction of reads originating from an already-­observed UMI. This is a function of library complexity and sequencing depth. More specifically, this is the fraction of confidently mapped, valid cell-­barcode, valid UMI reads that had a non-­unique (cell-­barcode, UMI, gene).")


def multiRunCommentsAndNotes1Sample(doc, values, data, runInfo, libraries):
    de = {i:data.parse(i)for i in data.sheet_names}
    de_sample = data.parse('Sample')

    rc = doc.paragraphs[16]
    rc.text = ""

    prev = False
    if any('Gene Expression' in i for i in de['Library']):
        rangemin = str(int(min(list(de['Library'][u'Gene Expression Number of reads']))/1000000))
        cellmin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Cells']))))
        cellmax = "{:,}".format(int(max(list(de['Sample'][u'Gene Expression Cells']))))
        readmin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Median reads per cell']))))
        readmax ="{:,}".format(int(max(list(de['Sample'][u'Gene Expression Median reads per cell']))))
        medgenemin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Median genes per cell']))))
        medgenemax ="{:,}".format(int(max(list(de['Sample'][u'Gene Expression Median genes per cell']))))
        totgenemin = "{:,}".format(int(min(list(de['Sample'][u'Gene Expression Total genes detected']))))
        totgenemax ="{:,}".format(int(max(list(de['Sample'][u'Gene Expression Total genes detected']))))

        runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
        runner = rc.add_run("Gene Expression")
        runner.bold = True
        if libraries == None:
            runner = rc.add_run(" library was sequenced on a ")
            runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
            runner.bold = True
            runner = rc.add_run(f" run. The sample has sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
            runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
            runner.bold = True
        else:
            with open(libraries, 'r') as f:
                flowcells = set([line.strip().split(',')[1] for line in f if 'Gene Expression' in line.strip().split(',')[-1]])
                runner = rc.add_run(" libraries were sequenced on ")
                platform = dict()
                platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                uniq = list(set(platforms))
                uniq.sort()
                runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                runner.bold = True
                runner = rc.add_run(f". The sample has sequencing yields of {rangemin} million reads. The sequencing run was setup as ")
                runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                runner.bold = True
        runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
        runner = rc.add_run(f"{values['version']}")
        runner.bold = True
        runner = rc.add_run(f" software using the default parameters. The number of cells captured is {cellmin}  and the mean number of reads per cell is {readmin}. Cells with extremely low number of UMI counts were filtered out.The median number of genes found per cell is {medgenemin} and the total number of genes detected is {totgenemin}")

        notesParagraph = doc.paragraphs[30]
        bulletEntry(notesParagraph, "Number of Cell", "The number of barcodes associated with cell-containing partitions.")
        bulletEntry(notesParagraph, "Median Reads per Cell", "Median number of read pairs sequenced from the cells assigned to this sample.")
        bulletEntry(notesParagraph, "Median Genes per Cell", "The median number of genes detected per cell assigned to this sample. Detection is defined as the presence of at least 1 UMI count.")
        bulletEntry(notesParagraph, "Median UMI Counts per Cell", "The median number of UMI counts per cell barcode.")
        bulletEntry(notesParagraph, "Reads Mapped Confidently to Transcriptome", "Fraction of reads that mapped to a unique gene in the transcriptome with a high mapping quality score as reported by the aligner.")
        bulletEntry(notesParagraph, "Sequencing Saturation", "The fraction of reads originating from an already-­observed UMI. This is a function of library complexity and sequencing depth. More specifically, this is the fraction of confidently mapped, valid cell-­barcode, valid UMI reads that had a non-­unique (cell-­barcode, UMI, gene).")

        prev = True

    if any('VDJ' in i for i in de['Library']):
        if prev:
            rc.add_run("\n\n")
        rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'VDJ' in i and 'Number of reads' in i])/1000000))
        cellmin = "{:,}".format(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'VDJ' in i and 'Estimated number of cells' in i])))
        cellmax = "{:,}".format(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'VDJ' in i and 'Estimated number of cells' in i])))
        readmin = "{:,}".format(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'VDJ' in i and 'Mean reads per cell' in i])))
        readmax = "{:,}".format(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'VDJ' in i and 'Mean reads per cell' in i])))
        vjspanmin = "{:,}".format(int(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'VDJ' in i and 'Number of cells with productive V-J spanning pair' in i])))
        vjspanmax = "{:,}".format(int(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'VDJ' in i and 'Number of cells with productive V-J spanning pair' in i])))
        clonomin = "{:,}".format(float(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'VDJ' in i and 'Paired clonotype diversity' in i])))
        clonomax = "{:,}".format(float(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'VDJ' in i and 'Paired clonotype diversity' in i])))
        runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
        runner = rc.add_run("VDJ")
        runner.bold = True
        if libraries == None:
            runner = rc.add_run(" library was sequenced on a ")
            runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
            runner.bold = True
            runner = rc.add_run(f" run. The sample has sequencing yield of {rangemin} million read. The sequencing run was setup as a ")
            runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
            runner.bold = True
        else:
            with open(libraries, 'r') as f:
                flowcells = set([line.strip().split(',')[1] for line in f if 'VDJ' in line.strip().split(',')[-1]])
                runner = rc.add_run(" library was sequenced on ")
                platform = dict()
                platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                uniq = list(set(platforms))
                uniq.sort()
                runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                runner.bold = True
                runner = rc.add_run(f". The sample has sequencing yield of {rangemin} million read. The sequencing run was setup as ")
                runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                runner.bold = True
        runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
        runner = rc.add_run(f"{values['version']}")
        runner.bold = True
        runner = rc.add_run(f" software using the default parameters. The number of cells captured is {cellmin} and mean number of reads per cell is {readmin}. Cells with extremely low number of UMI counts were filtered out. The number of cells with productive V-J spanning pair is {vjspanmin}. The paired clonotype diversity is {clonomin}.")

        notesParagraph = doc.paragraphs[30]
        bulletEntry(notesParagraph, "Number of Cells with Productive V-J Spanning Pair", "Fraction of cell-associated barcodes with at least one productive contig for each chain of the receptor pair.")
        bulletEntry(notesParagraph, "Reads Mapped to Any V(D)J Gene", "Fraction of reads that partially or wholly map to any germline V(D)J gene segment.")
        bulletEntry(notesParagraph, "Paired Clonotype Diversity", "Effective diversity of the paired clonotypes, computed as the Inverse Simpson Index of the clonotype frequencies.")
        prev = True

        if any('Antibody Capture' in i for i in de['Library']):
            if prev:
                rc.add_run("\n\n")
                rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Number of reads' in i])/1000000))
                cellmin = "{:,}".format(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Estimated number of cells' in i])))
                cellmax = "{:,}".format(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Estimated number of cells' in i])))
                antireadmin = "{:,}".format(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Mean reads per cell' in i])))
                antireadmax = "{:,}".format(int(max([max(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Mean reads per cell' in i])))
                antifracreadmin = "{:,}".format(float(min([min(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Fraction antibody reads usable' in i]))*100)
                antifracreadmax = "{:,}".format(float(max([max(list(de['Library'][i])) for i in de['Library'] if 'Antibody Capture' in i and 'Fraction antibody reads usable' in i]))*100)
                antifracreadusablemin = "{:,}".format(int(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Antibody reads usable per cell' in i])))
                antifracreadusablemax = "{:,}".format(int(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Antibody reads usable per cell' in i])))
                antiumimin = "{:,}".format(float(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
                antiumimax = "{:,}".format(float(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
                runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
                runner = rc.add_run("Antibody")
                runner.bold = True
                if libraries == None:
                    runner = rc.add_run(" libraries were sequenced on a ")
                    runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
                    runner.bold = True
                    runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
                    runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
                    runner.bold = True
                else:
                    with open(libraries, 'r') as f:
                        flowcells = set([line.strip().split(',')[1] for line in f if 'Antibody' in line.strip().split(',')[-1]])
                        runner = rc.add_run(" libraries were sequenced on ")
                        platform = dict()
                        platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                        uniq = list(set(platforms))
                        uniq.sort()
                        runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                        runner.bold = True
                        runner = rc.add_run(f". All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as ")
                        runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                        runner.bold = True
                runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
                runner = rc.add_run(f"{values['version']}")
                runner.bold = True
                runner = rc.add_run(f" software using the default parameters. The number of cells captured ranges from {cellmin} to {cellmax}. Cells with extremely low number of UMI counts were filtered out. The mean antibody reads per cell ranged from {antireadmin} to {antireadmax}. The fraction of antibody library reads that contain a recognized antibody barcode range from {antifracreadmin}% to {antifracreadmax}%. The number of antibody reads usable per cell ranges from {antifracreadusablemin} to {antifracreadusablemax}. The median antibody UMIs per cell is between {antiumimin} and {antiumimax}.")
                prev = True

    if any('Multiplexing Capture' in i for i in de['Library']):
        if prev:
            rc.add_run("\n\n")
            rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Number of reads' in i])/1000000))
            cellassignmin = "{:,}".format(int(min([min([j.split(' ')[0] for j in list(de['Library'][i])]) for i in de['Library'] if 'Multiplexing Capture' in i and 'Cells assigned to a sample' in i])))
            cellassignmax = "{:,}".format(int(max([max([j.split(' ')[0] for j in list(de['Library'][i])]) for i in de['Library'] if 'Multiplexing Capture' in i and 'Cells assigned to a sample' in i])))
            sampleassignmin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Samples assigned at least one cell' in i])))
            fracreadsmin = "{:}".format(float(min([min(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Fraction reads in cell-associated barcodes' in i]))*100)
            fracreadsmax = "{:,}".format(float(max([max(list(de['Library'][i])) for i in de['Library'] if 'Multiplexing Capture' in i and 'Fraction reads in cell-associated barcodes' in i]))*100)
            antifracreadusablemin = "{:,}".format(int(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Antibody reads usable per cell' in i])))
            antifracreadusablemax = "{:,}".format(int(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Antibody reads usable per cell' in i])))
            antiumimin = "{:,}".format(float(min([min(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
            antiumimax = "{:,}".format(float(max([max(list(de['Sample'][i])) for i in de['Sample'] if 'Antibody Capture' in i and 'Median UMI counts per cell' in i])))
            runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
            runner = rc.add_run("Multiplexing")
            runner.bold = True
            if libraries == None:
                runner = rc.add_run(" libraries were sequenced on a ")
                runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
                runner.bold = True
                runner = rc.add_run(f" run. All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as a ")
                runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run")
                runner.bold = True
            else:
                with open(libraries, 'r') as f:
                    flowcells = set([line.strip().split(',')[1] for line in f if 'Multiplexing' in line.strip().split(',')[-1]])
                    runner = rc.add_run(" libraries were sequenced on ")
                    platform = dict()
                    platforms = [runInfo[run]['platform'] for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']]
                    uniq = list(set(platforms))
                    uniq.sort()
                    runner = rc.add_run(f"{', '.join([f'a {i}' if platforms.count(i) == 1 else f'{platforms.count(i)} {i}' for i in uniq])} {['runs' if len(platforms) > 1 else 'run'][0]}")
                    runner.bold = True
                    runner = rc.add_run(f". All samples have sequencing yields of more than {rangemin} million read per sample. The sequencing run was setup as ")
                    runner = rc.add_run(', '.join([f"{run[0]} cycles + {run[1]} cycles" for run in set([(runInfo[run]['read1'], runInfo[run]['read2']) for i in flowcells for run in runInfo if i == runInfo[run]['flowcell']])]))
                    runner.bold = True
            runner = rc.add_run(f". Demultiplexing was done allowing 1 mismatch in the barcodes. The analysis was performed with the Cell Ranger ")
            runner = rc.add_run(f"{values['version']}")
            runner.bold = True
            runner = rc.add_run(f" software using the default parameters. The number of cells assigned to a sample ranges from {cellassignmin} to {cellassignmax}. There were a minimum of {sampleassignmin} samples assigned at least one cell. The fraction of reads in cell-associated barcodes range from {fracreadsmin}% to {fracreadsmax}%.")
            prev = True
    
    if any('CRISPR Guide Capture' in i for i in de['Library']):
        if prev:
            rc.add_run("\n\n")
            rangemin = str(int(min([min(list(de['Library'][i])) for i in de['Library'] if 'CRISPR Guide Capture' in i and 'Number of reads' in i])/1000000))
            meanpercellmin = "{:,}".format(int(min([min([j for j in list(de['Library'][i])]) for i in de['Library'] if 'CRISPR Guide Capture' in i and 'Mean reads per cell' in i])))
            pctprotospacermin = "{:.2%}".format(float(min([min([j for j in list(de['Sample'][i])]) for i in de['Sample'] if 'CRISPR Guide Capture' in i and 'with one or more protospacers detected' in i]))) 
            pct2protospacermin = "{:.2%}".format(float(min([min([j for j in list(de['Sample'][i])]) for i in de['Sample'] if 'CRISPR Guide Capture' in i and 'with two or more protospacers detected' in i]))) 
            runner = rc.add_run(f"{values['count']} 10x Genomics Single Cell ")
            runner = rc.add_run("CRISPR Guide Capture")
            runner.bold = True
            runner = rc.add_run(" libraries was sequenced on a ")
            runner = rc.add_run(f"{values['platform'].replace('_', ' ')}")
            runner.bold = True
            runner = rc.add_run(f" run. The sequencing yields is {rangemin} million reads. The sequencing run was setup as a ")
            runner = rc.add_run(f"{values['read1']} cycles + {values['read2']} cycles {values['sym']} run. ")
            runner.bold = True
            runner = rc.add_run(f"The mean reads per cell is {meanpercellmin}. The percentage of cells with one or more protospacers detected is {pctprotospacermin}. The percentage of cells with two or more protospacers detected is {pct2protospacermin}.")
            prev = True 
