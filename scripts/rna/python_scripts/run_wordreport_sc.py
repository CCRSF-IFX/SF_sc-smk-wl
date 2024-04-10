#!/usr/bin/env python2
# -*- coding: utf-8 -*-
#shent2 on 2018/09/13: modify the value of trueseq (library prep protocol) and add "check before delivery"
#shent2 on 2018/09/27: remove "check before delivery" and make "library prep protocol" bold in the first table
#chenv3 on 2020/09/11: updated to be compatible with cellranger 4.0.0
"""
Created on Tue Nov 21 23:58:30 2017

@author: Jack

uasge: python run_wordreport.py MikeBustin_ShaofeiZhang_CS022033_12Chiplib_110817_HVNYWBGX3.xlsx

"""
import os, sys, re
import pandas as pd
import time
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

import numpy as np

def replace_string(doc, before, after):
    for p in doc.paragraphs:
        if before in p.text:
            inline = p.runs

            for i in range(len(inline)):
                if before in inline[i].text:
                    text = inline[i].text.replace(before, after)
                    inline[i].text = text
                    #return text

    return 1


excelfile = sys.argv[1]
metafile = sys.argv[2]
cellranger = sys.argv[3]

filenm=os.path.splitext(metafile)[0][:-9]+'.docx'
emailfile = os.path.splitext(metafile)[0][:-9]+'_email.docx'
data = pd.ExcelFile(excelfile)

metadata = pd.read_csv(metafile, sep='\t')



if str(data.sheet_names[0]) != "metrics_summary":
    de = data.parse(data.sheet_names[0])
else:
    de = data.parse('metrics_summary')

#summary=data.parse('Summary')

if "RNA" in metadata['Application'][0].upper():
    docfile = '/is2/projects/CCR-SF/active/Software/scripts/lib/wordreporttemp/10X_sc_rna.docx'
    doc = Document(docfile)
else:
    docfile = '/is2/projects/CCR-SF/active/Software/scripts/lib/wordreporttemp/Yam_chip.docx'
    doc = Document(docfile)


pi = metadata[metadata.columns[1]][0]
lc = metadata[metadata.columns[2]][0]
bc = metadata[metadata.columns[4]][0]
pt = metadata[metadata.columns[3]][0]
csac = str(metadata[metadata.columns[10]][0])
stip = str(len(metadata))
sitr = str(len(de))
date = time.strftime("%x")

replace_string(doc,"{PI}", pi )
replace_string(doc,"{LC}", lc)
replace_string(doc,"{BC}", bc )
replace_string(doc,"{PT}", pt)
replace_string(doc,"{CSAC}", csac )
replace_string(doc,"{STIP}", stip)
replace_string(doc,"{SITR}", sitr )
replace_string(doc,"{DATE}", date)

seqType = ""
seqType = metadata['Application'][0]

pattern = re.compile(r'\d+\.\d+\.\d+')
if len(pattern.findall(cellranger)) == 1:
    version = pattern.findall(cellranger)[0]
else:
    version = "UNKNOWN"

#st = str(len(summary)) + ' ' + str(metadata[metadata.columns[8]][0])
st = str(len(de)) + ' ' + seqType
mt = str(metadata[metadata.columns[11]][0])

if 'x' in metadata['ReadLength'][0]:
    if '2x' in metadata['ReadLength']:
        sym = 'symmetric'
        r1 = metadata['ReadLength'][0].split('x')[1]
        r2 = r1
    else:
        sym = 'non-symmetric'
        r1 = metadata['ReadLength'][0].split('x')[1]
        r2 = '0'
else:
	temp = re.findall(": *(\d+)", metadata['ReadLength'][0])
	r1 = temp[0]
	r2 = temp[1]
	if r1 == r2:
	    sym = 'symmetric'
	else:
	    sym = 'non-symmetric'
#if int(metadata['ReadLength'][0].split('x')[0]) == 2 or len(metadata['ReadLength'][0].split(',')) == 2:
#    ps = 'paired-end'
#elif int(metadata['ReadLength'][0].split('x')[0]) == 1:
#    ps = 'single-end'
#else:
#    ps = ""

barcode = str(float(min(list(de[u'Q30 Bases in Barcode'])))*100)
#index = str(float(min(list(de[u'Q30 Bases in Sample Index'])))*100)
umi = str(float(min(list(de[u'Q30 Bases in UMI'])))*100)
cellmin = "{:,}".format(int(min(list(de[u'Estimated Number of Cells']))))
cellmax = "{:,}".format(int(max(list(de[u'Estimated Number of Cells']))))

if sum(['RNA Read' in i for i in de.columns ]) == 1:
    rnaread = str(float(min(list(de[u'Q30 Bases in RNA Read'])))*100)
else:
    rnaread = str(min([float(min(list(de[u'Q30 Bases in RNA Read']))), float(min(list(de[u'Q30 Bases in RNA Read 2'])))])*100)

#if "RNA" in metadata['Application'][0].upper():
rangemin = str(int(min(list(de[u'Number of Reads']))/1000000))
rangemax = str(int(max(list(de[u'Number of Reads']))/1000000))
readmin = "{:,}".format(int(min(list(de[u'Mean Reads per Cell']))))
readmax ="{:,}".format(int(max(list(de[u'Mean Reads per Cell']))))
medgenemin = "{:,}".format(int(min(list(de[u'Median Genes per Cell']))))
medgenemax ="{:,}".format(int(max(list(de[u'Median Genes per Cell']))))
totgenemin = "{:,}".format(int(min(list(de[u'Total Genes Detected']))))
totgenemax ="{:,}".format(int(max(list(de[u'Total Genes Detected']))))
# elif "VDJ" in metadata['Application'][0].upper():
#     rangemin = str(int(min(list(de[u'Number of Read Pairs']))/1000000))
#     rangemax = str(int(max(list(de[u'Number of Read Pairs']))/1000000))
#     readmin = str(int(min(list(de[u'Mean Read Pairs per Cell']))))
#     readmax =str(int(max(list(de[u'Mean Read Pairs per Cell']))))

rc = doc.paragraphs[13]

replace_string(doc, "ST", str(len(de)))
replace_string(doc, "MT", mt.replace('_', ' '))
replace_string(doc, "R1", r1)
replace_string(doc, "R2", r2)
replace_string(doc, "SYM", sym)
replace_string(doc, "BARCODEQ30", barcode)
#replace_string(doc, "INDEXQ30", index)
replace_string(doc, "UMIQ30", umi)
replace_string(doc, "CELLMIN", cellmin)
replace_string(doc, "CELLMAX", cellmax)
replace_string(doc, "RNAREADQ30", rnaread)
replace_string(doc, "RANGEMIN", rangemin)
replace_string(doc, "RANGEMAX", rangemax)
replace_string(doc, "READMIN", readmin)
replace_string(doc, "READMAX", readmax)
replace_string(doc, "VERSION", version)

if "RNA" in metadata['Application'][0].upper():
    replace_string(doc, "MEDGENEMIN", medgenemin)
    replace_string(doc, "MEDGENEMAX", medgenemax)
    replace_string(doc, "TOTGENEMIN", totgenemin)
    replace_string(doc, "TOTGENEMAX", totgenemax)

for i in rc.runs:
    i.text = i.text.replace("{", "")
    i.text = i.text.replace("}", "")

#if "VDJ" in metadata['Application'][0].upper():

flowcell = str(metafile.split('_')[-2])

table1 = doc.tables[0]
table1.cell(0,2).text = flowcell
table1.cell(0,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(0,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(1,2).text = str(metadata[metadata.columns[11]][0])
table1.cell(1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(1,2).paragraphs[0].paragraph_format.space_after = 0
#table1.cell(2,2).text = str(metadata[metadata.columns[8]][0])
table1.cell(2,2).text = seqType
table1.cell(2,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(2,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(3,2).text = str(metadata[metadata.columns[12]][0])
table1.cell(3,2).paragraphs[0].runs[0].font.bold = True
table1.cell(3,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(3,2).paragraphs[0].paragraph_format.space_after = 0
table1.cell(4,2).text = str(metadata[metadata.columns[13]][0]).split('=')[1] + ' samples in 1 run'
table1.cell(4,2).paragraphs[0].runs[0].font.bold = True
table1.cell(4,2).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(4,2).paragraphs[0].paragraph_format.space_after = 0
#table1.cell(4,5).text = metadata[metadata.columns[9]][0]

if "RNA" in metadata['Application'][0].upper():
    if "3'" in metadata['Application'][0].upper():
        trueseq = "Chromium 3' RNA-seq library prep"
        library = "3' mRNA"
    elif "5'" in metadata['Application'][0].upper():
        trueseq = "Chromium 5' RNA-seq library prep"
        library = "5' mRNA"
    else:
        trueseq = "Chromium RNA-seq library prep"
        library = "mRNA"
elif "VDJ" in metadata['Application'][0].upper():
    trueseq = "Chromium V(D)J Enrichment Kit"
    library = "VDJ"
else:
    trueseq = "UNKNOWN library preparation protocol"
    library = "UNKNOWN"

replace_string(doc, "TRUESEQ", library)

table1.cell(2,5).text = str(trueseq)
table1.cell(2,5).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(2,5).paragraphs[0].paragraph_format.space_after = 0
table1.cell(2,5).paragraphs[0].runs[0].font.bold = True

if "HG38" in metadata['ReferenceGenome'][0].upper() and "RNA" in metadata['Application'][0].upper():
    table1.cell(4,5).text = "hg38 - 2020-A"
elif "MM10" in metadata['ReferenceGenome'][0].upper() and "RNA" in metadata['Application'][0].upper():
    table1.cell(4,5).text = "mm10 - 2020-A"
else:
    table1.cell(4,5).text = "UNKNOWN"
table1.cell(4,5).paragraphs[0].runs[0].font.bold = True
table1.cell(4,5).paragraphs[0].paragraph_format.space_before = Pt(2)
table1.cell(4,5).paragraphs[0].paragraph_format.space_after = 0

p = doc.tables[3]
p.cell(3,1).text = p.cell(3,1).text.replace("{VERSION}", version)
p.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
p.cell(3,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p=doc.tables[4]
p.style = 'Table Grid'
for i in range(len(de)):
    p.add_row()
    p.cell(i+1,0).text = str(de[u"Sample"][i])
    p.cell(i+1,0).paragraphs[0].paragraph_format.space_before = Pt(2)
    p.cell(i+1,0).paragraphs[0].paragraph_format.space_after = 0
    p.cell(i+1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.cell(i+1,1).text = "{:,}".format(int(de[u"Estimated Number of Cells"][i]))
    p.cell(i+1,1).paragraphs[0].paragraph_format.space_before = Pt(2)
    p.cell(i+1,1).paragraphs[0].paragraph_format.space_after = 0
    p.cell(i+1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.cell(i+1,2).text = "{:,}".format(int(de[u"Mean Reads per Cell"][i]))
    p.cell(i+1,2).paragraphs[0].paragraph_format.space_before = Pt(2)
    p.cell(i+1,2).paragraphs[0].paragraph_format.space_after = 0
    p.cell(i+1,2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.cell(i+1,3).text = "{:,}".format(int(de[u"Median Genes per Cell"][i]))
    p.cell(i+1,3).paragraphs[0].paragraph_format.space_before = Pt(2)
    p.cell(i+1,3).paragraphs[0].paragraph_format.space_after = 0
    p.cell(i+1,3).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.cell(i+1,4).text = "{:,}".format(int(de[u"Total Genes Detected"][i]))
    p.cell(i+1,4).paragraphs[0].paragraph_format.space_before = Pt(2)
    p.cell(i+1,4).paragraphs[0].paragraph_format.space_after = 0
    p.cell(i+1,4).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.cell(i+1,5).text = "{:,}".format(int(de[u"Median UMI Counts per Cell"][i]))
    p.cell(i+1,5).paragraphs[0].paragraph_format.space_before = Pt(2)
    p.cell(i+1,5).paragraphs[0].paragraph_format.space_after = 0
    p.cell(i+1,5).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT


doc.save(filenm)

docfile = '/is2/projects/CCR-SF/active/Software/scripts/lib/wordreporttemp/ccremail.docx'
doc = Document(docfile)
app = str(metadata['Application'][0])
pilast = pi.replace("_", " ").split()[1]
ttext = rc.text

replace_string(doc,"{App}", app)
replace_string(doc,"{Flowcell1}", flowcell)
replace_string(doc,"Pilast", pilast)
replace_string(doc,"{ST}", st)
replace_string(doc,"{CSAC}", csac)
replace_string(doc,"{Flowcell2}", flowcell)

doc.paragraphs[9].text = rc.text
rows,columns = de.shape
tb=doc.tables[0].rows[0].cells[0].add_table(1, columns)
tb.style = 'Table Grid'

h = list(de)
head = tb.rows[0].cells
for idx, name in enumerate(h):
    paragraph = head[idx].paragraphs[0]
    run = paragraph.add_run(name)
    run.bold = True

for r in range(rows):
    cells = tb.add_row().cells
    for col in range(columns):
        if str(de.iat[r,col]).count('.') == 1:
            print("{0:.3f}".format(de.iat[r,col]))
            cells[col].text = "{0:.3f}".format(de.iat[r,col])
        else:
            cells[col].text = str(de.iat[r,col])

doc.save(emailfile)


out = ""
out += os.getcwd().split("/")[-1] + ","
out += os.path.basename(excelfile).split('.')[0] + ","
adate = time.strftime('%m/%d/%y', time.gmtime(os.path.getctime(excelfile)))
out += str(adate) + ",cronjob,," + str(adate) + ","
#if str(de[de.columns[2]][1]) == "Yield (MBases)":
#    gb = str(float(str(de[de.columns[2]][2]).replace(",", ""))/1000)
#    out += gb + ","
#    out += str(len(summary)) + ","
out += str(metadata['Application'][0]) + ",,"
out += str(adate) + "\n"

print(out)
#csvrecord = "/is2/projects/CCR-SF/scratch/illumina/Processing/CCRSFIFX_ANALYSIS/processed.csv"
#ck = open (csvrecord, 'r')
#ckk = ck.read()
#ck.close()
#if ckk.find(os.path.basename(excelfile).split('.')[0]) > 0:
#    pass
#else:
#    cr = open (csvrecord, 'a')
#    cr.write(out)
#    cr.close()
